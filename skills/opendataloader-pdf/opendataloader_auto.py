#!/usr/bin/env python3
"""opendataloader_auto.py v2.3.0 (2026-04-25)
表格结构提取: qwen2.5vl OCR + cell_bbox(坐标), col_span/row_span, 嵌套表头检测, 斜线表头检测, has_merged_cells标记。
文档自动检测与转换脚本(统一处理 PDF / Word / Excel)
版本: 2.0.0
日期: 2026-04-24

更新说明:
  - 扫描 PDF 的处理优先级调整:qwen2.5vl(主)→ Hybrid/EasyOCR(备)→ Fast(保底)
  - 移除 `--force-qwen` 标志,扫描 PDF 自动优先使用 qwen2.5vl
  - `--force-qwen` 重新定义为:跳过 Hybrid,直接用 qwen + Fast(禁用 Hybrid 作为备选)

架构说明(v1.8):
  输入为 PDF  → 直接提取(位置信息 + 内容,PDF 限制:有 bbox/页码,表格数据可能稀疏)
  输入为 Word → Word转PDF + docx直接提取 + PDF位置提取 → 智能合并
               相同信息以更可靠来源为准:表格/文本内容以 docx 为准,位置信息(页码/bbox)以 PDF 为准
  输入为 Excel → openpyxl直接提取(结构化数据/表格/图表)+ Excel转PDF获取页码 → 合并

输出:统一的两文件架构
  ✦ {basename}.json    ← 结构化数据,含完整内容和位置信息
  ✦ {basename}.md      ← 人类可读的 Markdown
"""

import sys
import os
import json
import argparse
import subprocess
import time
import re
from pathlib import Path
from docx.oxml.ns import qn as _qn


# ---------- 辅助函数 ----------
def _extract_excel_table_with_spans(ws) -> list:
    """
    从 openpyxl Worksheet 中提取带 colspan/row_span 的表格数据。
    返回格式(与 Word / qwen2.5vl 输出一致):
    [
        {
            "cells": ["列1", "列2", "列3"],
            "col_span": [1, 1, 1],
            "row_span": [1, 1, 1],
        },
        ...
    ]

    实现逻辑:
    1. 遍历所有合并单元格范围,建立 {cell_coord: (span_cols, span_rows, is_merged_start)} 映射
    2. 逐行逐列读取,若单元格属于合并范围但非起始格,记为 ''(空内容)
    3. 合并范围的起始格,记录实际的 colspan/row_span
    4. 横向合并:colspan > 1
    5. 纵向合并:row_span > 1
    """
    from openpyxl.utils import range_boundaries

    # 步骤1:建立合并单元格映射
    # merged_ranges: list of (min_col, min_row, max_col, max_row)
    merged_map = {}  # (row, col) -> {colspan, rowspan, is_start} or {skip: True}
    for merged_range in ws.merged_cells.ranges:
        min_col, min_row, max_col, max_row = range_boundaries(str(merged_range))
        span_cols = max_col - min_col + 1
        span_rows = max_row - min_row + 1
        # 标记起始单元格
        merged_map[(min_row, min_col)] = {"colspan": span_cols, "rowspan": span_rows, "is_start": True}
        # 标记被合并的后续单元格(跳过)
        for r in range(min_row, max_row + 1):
            for c in range(min_col, max_col + 1):
                if (r, c) != (min_row, min_col):
                    merged_map[(r, c)] = {"skip": True, "is_merged": True}

    # 步骤2:逐行逐列输出,跳过被合并的单元格
    result = []
    for row in ws.iter_rows(values_only=False):
        row_cells_out = []
        col_spans = []
        row_spans = []
        col_ptr = 0
        row_num = row[0].row  # 当前行号(1-indexed)

        while col_ptr < len(row):
            cell = row[col_ptr]
            r, c = cell.row, cell.column  # 1-indexed
            info = merged_map.get((r, c), {})

            if info.get("skip"):
                # 被合并的格子(横向/纵向都属于起始格覆盖范围),跳过
                col_ptr += 1
                continue

            if info.get("is_start"):
                cs = info["colspan"]
                rs = info["rowspan"]
            else:
                cs = 1
                rs = 1

            val = cell.value
            cell_text = _make_serializable(val) if val is not None else ""
            row_cells_out.append(cell_text)
            col_spans.append(cs)
            row_spans.append(rs)
            col_ptr += 1

        # 整行被纵向合并覆盖(无任何内容输出),跳过
        if not row_cells_out:
            continue

        result.append({
            "cells": row_cells_out,
            "col_span": col_spans,
            "row_span": row_spans,
        })

    return result


def _parse_markdown_table(md_text: str) -> list:
    """
    解析 Markdown 表格文本为 2D 数组。
    输入: "列1 | 列2 | 列3\n值1 | 值2 | 值3\n..."
    输出: [["列1", "列2", "列3"], ["值1", "值2", "值3"], ...]
    """
    rows = []
    for line in md_text.strip().split('\n'):
        line = line.strip()
        if not line or line.startswith('|---') or line.startswith('---'):
            continue
        # 去掉首尾的 |
        if line.startswith('|'):
            line = line[1:]
        if line.endswith('|'):
            line = line[:-1]
        cells = [c.strip() for c in line.split('|')]
        rows.append(cells)
    return rows


def _extract_docx_table_with_spans(table) -> list:
    """
    从 python-docx Table 中提取带 colspan/row_span 的表格数据。
    返回格式(与 qwen2.5vl OCR 输出一致):
    [
        {
            "cells": ["A", "B", "C"],
            "col_span": [1, 2, 1],   # 每列的 colspan
            "row_span": [1, 1, 1],   # 每列的 row_span(垂直合并)
        },
        ...
    ]
    检测:
    - colspan: gridSpan 元素(Word 单元格横向合并)
    - rowspan: vMerge 元素(Word 单元格纵向合并)
    """
    from docx.oxml.ns import qn

    result = []
    for row in table.rows:
        cells = row.cells
        col_spans = []
        row_spans = []
        cell_texts = []

        for cell in cells:
            tc = cell._tc
            # colspan: gridSpan
            gridSpan_el = tc.find(_qn('w:gridSpan'))
            cs = int(gridSpan_el.get(_qn('w:val'))) if gridSpan_el is not None else 1
            col_spans.append(cs)

            # rowspan: vMerge(continue=续行,None/restart=起始行)
            vMerge_el = tc.find(_qn('w:vMerge'))
            if vMerge_el is not None:
                val = vMerge_el.get(_qn('w:val'))
                rs = -1 if val == 'continue' else 1  # -1 表示属于上面的纵向合并
            else:
                rs = 1
            row_spans.append(rs)

            cell_texts.append(cell.text.strip())

        result.append({
            "cells": cell_texts,
            "col_span": col_spans,
            "row_span": row_spans,
        })

    return result


def _detect_merged_cells_from_flat(table_data: list) -> tuple:
    """
    从 opendataloader 输出的扁平 table_data 推算是否存在合并单元格。
    table_data: 扁平 2D 数组 [[cell, ...], ...]
    返回: (flat_array, has_merged_cells: bool)

    推算逻辑:
    - 如果某行 cell 数 < 最大列数 → 存在横向合并(colspan)
    - 如果相邻两行某列内容完全相同且其间无空行 → 可能纵向合并(row_span)
    缺点:无法精确还原 colspan/row_span 值,只能标记"疑似合并"
    """
    if not table_data or not isinstance(table_data, list):
        return table_data, False

    # 估算总列数(最多 cell 的那一行)
    max_cols = max(len(row) for row in table_data if isinstance(row, list))
    if max_cols == 0:
        return table_data, False

    has_merged = False
    for row in table_data:
        if not isinstance(row, list):
            continue
        cell_count = len([c for c in row if c is not None and str(c).strip() != ""])
        if len(row) < max_cols:
            has_merged = True
            break

    return table_data, has_merged


def _table_to_markdown_with_spans(table_rows: list, output_format: str = "docx") -> str:
    """
    将带 colspan/row_span 的表格数据渲染为 Markdown。
    output_format="docx": 使用 docx 的 gridSpan 渲染
    output_format="inferred": 使用推算的合并信息渲染
    横向合并的单元格:在对应位置填满空列
    纵向合并的单元格:在续行位置填入 "^"(Markdown table rowspan 扩展惯例)
    """
    if not table_rows:
        return ""

    # 推算总列数
    def total_cols(row_entry):
        if isinstance(row_entry, dict):
            return sum(row_entry.get("col_span", [1] * len(row_entry["cells"])))
        return len(row_entry)

    # 扁平渲染一行
    def render_flat_row(row_entry):
        if isinstance(row_entry, dict):
            cells = row_entry["cells"]
            col_spans = row_entry.get("col_span", [1] * len(cells))
            row_spans = row_entry.get("row_span", [1] * len(cells))
            rendered = []
            for i, cell in enumerate(cells):
                cs = col_spans[i] if i < len(col_spans) else 1
                rendered.append(cell if cell else "")
                for _ in range(cs - 1):
                    rendered.append("")  # 横向合并的空列位
            return rendered
        return list(row_entry)

    max_cols = max(total_cols(r) for r in table_rows) if table_rows else 0
    if max_cols == 0:
        return ""

    lines = []
    # 表头行
    header = render_flat_row(table_rows[0])
    while len(header) < max_cols:
        header.append("")
    lines.append("| " + " | ".join(str(c) for c in header[:max_cols]) + " |")
    # 分隔行
    lines.append("| " + " | ".join(["---"] * max_cols) + " |")
    # 数据行
    for row_entry in table_rows[1:]:
        flat = render_flat_row(row_entry)
        while len(flat) < max_cols:
            flat.append("")
        lines.append("| " + " | ".join(str(c) for c in flat[:max_cols]) + " |")

    return "\n".join(lines)


def _parse_table_data(raw_data):
    """
    统一解析 table_data 字段,支持两种格式:
    - 简单格式: [["列1","列2"], ["值1","值2"]]       → 直接返回,has_span=False
    - 增强格式: [{"cells":[...],"col_span":[...],"row_span":[...]}] → 展开为 2D 数组

    返回: (flat_2d_array, has_merged_cells: bool)
    has_merged_cells=True 表示该表格存在跨行/跨列单元格
    """
    if not raw_data or not isinstance(raw_data, list):
        return [], False
    if raw_data and isinstance(raw_data[0], list):
        return raw_data, False
    try:
        total_cols = 0
        for row_obj in raw_data:
            cells = row_obj.get("cells", [])
            col_span = row_obj.get("col_span", [])
            if col_span and len(col_span) == len(cells):
                total_cols = max(total_cols, sum(col_span))
            else:
                total_cols = max(total_cols, len(cells))
        flat = []
        has_span = False
        for row_obj in raw_data:
            cells = row_obj.get("cells", [])
            col_sp = row_obj.get("col_span", [1] * len(cells))
            row_sp = row_obj.get("row_span", [1] * len(cells))
            flat_row = []
            col_ptr = 0
            for i, cell in enumerate(cells):
                cs = col_sp[i] if i < len(col_sp) else 1
                rs = row_sp[i] if i < len(row_sp) else 1
                if cs > 1 or rs > 1:
                    has_span = True
                flat_row.append(cell)
                for _ in range(cs - 1):
                    flat_row.append("")
                col_ptr += cs
            while len(flat_row) < total_cols:
                flat_row.append("")
            flat.append(flat_row)
        return flat, has_span
    except Exception:
        return [], False


def _analyze_table_structure(raw_td):
    """
    基于 cell_bbox 和 col_span/row_span 分析表格结构。
    返回:
      header_depth: int         # 表头行数
      has_nested_headers: bool   # 是否存在嵌套表头
      has_diagonal_header: bool   # 是否存在斜线表头
      col_groups: list         # 列分组 [[父col_idx, [子col_idx...]], ...]
    """
    result = {
        "header_depth": 1,
        "has_nested_headers": False,
        "has_diagonal_header": False,
        "col_groups": [],
    }
    if not raw_td or not isinstance(raw_td, list):
        return result

    bboxes = []
    for row in raw_td:
        if isinstance(row, dict) and row.get("cell_bbox"):
            bboxes.append(row["cell_bbox"])
        else:
            bboxes.append(None)

    # 估算表格总列数
    total_cols = 0
    for ri, row in enumerate(raw_td):
        if not isinstance(row, dict):
            continue
        cells = row.get("cells", [])
        cs = row.get("col_span", [])
        if cs and len(cs) == len(cells):
            total_cols = max(total_cols, sum(cs))
        else:
            total_cols = max(total_cols, len(cells))

    # 斜线表头检测
    for ri, row in enumerate(raw_td):
        if not isinstance(row, dict):
            continue
        cells = row.get("cells", [])
        bb = bboxes[ri] if ri < len(bboxes) else None
        if not bb or len(cells) < 2:
            continue
        non_empty = [(i, c, bb[i]) for i, c in enumerate(cells)
                     if c and isinstance(bb[i], (list, tuple)) and len(bb[i]) >= 4]
        if len(non_empty) < 2:
            continue
        for k in range(len(non_empty) - 1):
            i1, c1, b1 = non_empty[k]
            i2, c2, b2 = non_empty[k + 1]
            y1b, y1t = b1[1], b1[3]
            y2b, y2t = b2[1], b2[3]
            if abs(y1b - y2b) > 5 and abs(y1t - y2t) < 10:
                result["has_diagonal_header"] = True
                break

    # 嵌套表头检测
    header_depth = 1
    for ri, row in enumerate(raw_td):
        if not isinstance(row, dict):
            continue
        rs = row.get("row_span", [])
        cells = row.get("cells", [])
        if not cells:
            continue
        if rs and all(r > 1 for r in rs):
            header_depth = max(header_depth, ri + 1)

    # 通过 col_span 推断列分组
    col_groups = []
    span_col_map = {}
    for ri, row in enumerate(raw_td):
        if not isinstance(row, dict):
            continue
        cells = row.get("cells", [])
        cs = row.get("col_span", [1] * len(cells))
        col_ptr = 0
        for i, cell in enumerate(cells):
            cs_val = cs[i] if i < len(cs) else 1
            if cs_val > 1:
                span_col_map[col_ptr] = {"row": ri, "cell": cell, "col_span": cs_val, "children": []}
            col_ptr += cs_val

    for start_col, info in span_col_map.items():
        span_len = info["col_span"]
        children = []
        for ri2, row2 in enumerate(raw_td):
            if not isinstance(row2, dict) or ri2 <= info["row"]:
                continue
            cells2 = row2.get("cells", [])
            cs2 = row2.get("col_span", [1] * len(cells2))
            col_ptr2 = 0
            for i2, cell2 in enumerate(cells2):
                rcs_val = cs2[i2] if i2 < len(cs2) else 1
                if col_ptr2 >= start_col and col_ptr2 < start_col + span_len and rcs_val == 1 and cell2:
                    children.append(col_ptr2)
                col_ptr2 += rcs_val
        info["children"] = sorted(set(children))
        if children or info["col_span"] > 1:
            col_groups.append([start_col, info["children"]])

    if header_depth > 1 or col_groups:
        result["has_nested_headers"] = True
    result["col_groups"] = col_groups
    result["header_depth"] = max(header_depth, 1)
    return result




def _render_pdf_pages_as_images(pdf_path: str, page_numbers: list,
                                  output_dir: str) -> dict:
    """
    将 PDF 指定页渲染为图像文件(用于 qwen2.5vl OCR 输入)。
    使用 pypdfium2 逐页渲染,返回 {page_num: image_path} 映射。
    """
    import pypdfium2 as pdfium
    page_images = {}
    os.makedirs(output_dir, exist_ok=True)
    try:
        pdf_doc = pdfium.PdfDocument(pdf_path)
        for pg in page_numbers:
            if pg < 1 or pg > len(pdf_doc):
                continue
            page = pdf_doc[pg - 1]  # 0-indexed
            bitmap = page.render(
                scale=2.0,  # 2x scale for better OCR quality
                stream=pdfium.PdfRenderActivity.RETURN
            )
            pil_img = bitmap.to_pil()
            img_path = os.path.join(output_dir, f"page_{pg:04d}.png")
            pil_img.save(img_path, format="PNG")
            page_images[pg] = img_path
    except Exception as e:
        print(f"[警告] 页面渲染失败: {e}", file=sys.stderr)
    return page_images


def _ocr_images_with_qwen(image_paths: list, lang: str = "auto") -> list:
    """
    将图像列表交给 qwen2.5vl OCR,返回元素列表(与 qwen_ocr_pdf 结构一致)。
    image_paths: [{page_num: image_path}, ...] 或简单列表
    """
    import urllib.request
    import base64
    import io
    from PIL import Image

    lang_prompt = "中文" if lang == "zh" else "英文"
    all_elements = []

    for img_path in image_paths:
        if not os.path.exists(img_path):
            continue
        try:
            with open(img_path, "rb") as f:
                img_data = base64.b64encode(f.read()).decode("utf-8")
            with Image.open(img_path) as pil_img:
                width, height = pil_img.size

            prompt = QWEN_OCR_PROMPT

            payload = {
                "model": QWEN_MODEL,
                "messages": [{
                    "role": "user",
                    "content": [
                        {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{img_data}"}},
                        {"type": "text", "text": prompt}
                    ]
                }],
                "stream": False
            }

            req = urllib.request.Request(
                OLLAMA_URL + "/chat",
                data=json.dumps(payload).encode("utf-8"),
                headers={"Content-Type": "application/json"},
                method="POST"
            )
            with urllib.request.urlopen(req, timeout=120) as resp:
                result = json.loads(resp.read().decode("utf-8"))

            content = result.get("message", {}).get("content", "")
            # 去掉 markdown code block
            content = re.sub(r"^```(?:json)?\s*", "", content.strip())
            content = re.sub(r"\s*```$", "", content.strip())
            data = json.loads(content)

            for block in data.get("blocks", []):
                btype = block.get("type", "")
                bcontent = (block.get("content") or "").strip()
                if not bcontent:
                    continue
                if btype == "text":
                    all_elements.append({
                        "type": "paragraph",
                        "content": bcontent,
                        "heading": block.get("heading", ""),
                    })
                elif btype == "table":
                    raw_td = block.get("table_data", [])
                    table_data, has_span = _parse_table_data(raw_td)
                    has_merged = any(
                        any(cs > 1 or rs > 1 for cs, rs in zip(
                            r.get("col_span", [1]*len(r.get("cells",[]))),
                            r.get("row_span", [1]*len(r.get("cells",[])))))
                        for r in raw_td
                    ) if raw_td else False
                    all_elements.append({
                        "type": "table",
                        "content": bcontent,
                        "table_data": table_data,
                        "has_merged_cells": has_span or has_merged,
                        "_raw_td": raw_td,
                    })
        except Exception as e:
            print(f"[警告] qwen2.5vl OCR 失败 {img_path}: {e}", file=sys.stderr)
    return all_elements


def _process_pages_with_qwen(pdf_path: str, output_dir: str,
                               table_pages: list, lang: str) -> list:
    """
    仅对指定页码列表渲染图像并执行 qwen2.5vl OCR。
    返回该批页的 kids 元素列表(与 qwen_ocr_pdf 结构一致)。
    """
    _got_vram = False
    _vm = None
    try:
        sys.path.insert(0, str(Path("/home/wangyc/.openclaw/scripts")))
        from vram_manager import VMgr
        _vm = VMgr()
        import time as _time
        _waited = 0
        while _waited < 300:
            _status = _vm.status()
            _vram_free = _status.get("vram_free_mb", 0)
            _state = _status.get("state", "?")
            if _vram_free >= 18000:
                if _vm.acquire_for_comfy(reason="opendataloader-qwen-ocr"):
                    _got_vram = True
                    break
            if _waited == 0:
                print(f"[      ] VRAM 排队等待(free={_vram_free}MB)...")
            _time.sleep(10)
            _waited += 10
    except Exception as _e:
        print(f"[警告] VRAM 队列异常: {_e}", file=sys.stderr)

    try:
        # 渲染指定页为图像
        render_dir = os.path.join(output_dir, "_table_pages_render")
        page_images = _render_pdf_pages_as_images(pdf_path, table_pages, render_dir)
        if not page_images:
            print(f"[警告] 页面渲染失败,返回空", file=sys.stderr)
            return []

        # 调用 qwen2.5vl OCR
        img_list = [img_path for pg, img_path in sorted(page_images.items())]
        ocr_elements = _ocr_images_with_qwen(img_list, lang=lang)

        # 建立页码映射(image_paths 顺序与 table_pages 顺序对应)
        pg_elements = {}
        for i, img_path in enumerate(img_list):
            pg_num = sorted(page_images.keys())[i]
            for elem in ocr_elements:
                e = dict(elem)
                e["page_number"] = pg_num
                if pg_num not in pg_elements:
                    pg_elements[pg_num] = []
                pg_elements[pg_num].append(e)

        # 收集所有 kids
        all_kids = []
        for pg in sorted(pg_elements.keys()):
            for elem in pg_elements[pg]:
                all_kids.append(elem)
        print(f"[      ] qwen2.5vl OCR 完成:{len(table_pages)} 页 → {len(all_kids)} 元素")
        return all_kids
    finally:
        if _vm is not None:
            try:
                _vm.release_and_restore()
            except Exception:
                pass


def _merge_consecutive_tables(elements: list) -> list:
    """
    检测并合并跨页表格。
    规则:
      - 连续两页出现相似结构的表格(列数相近或表头前缀相同)→ 合并
      - 合并后表格保留 continued_from / continued_on_page 标记
      - _segments 记录每段来源页码,方便回溯
    """
    if not elements:
        return elements
    result = []
    i = 0
    while i < len(elements):
        elem = elements[i]
        if elem.get("type") != "table" or elem.get("ocr_confidence") == "low":
            result.append(elem)
            i += 1
            continue
        page_num = elem.get("page_number", 0)
        table_data = elem.get("table_data", [])
        content_preview = (elem.get("content") or "")[:60].strip()
        if i + 1 < len(elements):
            nxt = elements[i + 1]
            if (nxt.get("type") == "table"
                    and nxt.get("page_number", 0) == page_num + 1
                    and nxt.get("ocr_confidence") != "low"):
                nxt_data = nxt.get("table_data", [])
                col_diff = abs(len(table_data[0] if table_data else [])
                              - len(nxt_data[0] if nxt_data else [])) if table_data and nxt_data else 999
                nxt_preview = (nxt.get("content") or "")[:60].strip()
                similar = (col_diff <= 2) or (
                    content_preview[:20] == nxt_preview[:20]
                    if content_preview and nxt_preview else False
                )
                if similar and nxt_data:
                    merged_data = list(table_data) + list(nxt_data)
                    merged_elem = dict(elem)
                    merged_elem["table_data"] = merged_data
                    merged_elem["content"] = (elem.get("content") or "") + "\n" + (nxt.get("content") or "")
                    merged_elem["continued_from"] = None
                    merged_elem["continued_on_page"] = nxt.get("page_number")
                    merged_elem["_segments"] = [
                        {"page": page_num, "table_data": list(table_data)},
                        {"page": nxt.get("page_number"), "table_data": list(nxt_data)}
                    ]
                    result.append(merged_elem)
                    i += 2
                    continue
        result.append(elem)
        i += 1
    return result


# ---------- Java 路径 ----------
JAVA_BIN = "/home/wangyc/opt/jre/amazon-corretto-11.0.30.7.1-linux-x64/bin/java"
JAVA_HOME = "/home/wangyc/opt/jre/amazon-corretto-11.0.30.7.1-linux-x64"

HYBRID_SERVER_PID_FILE = "/tmp/opendataloader_hybrid_server.pid"
HYBRID_SERVER_PORT = 5002

# ---------- qwen2.5vl OCR 配置(Hybrid 失败时的备用) ----------
QWEN_MODEL = "qwen2.5vl:latest"
OLLAMA_URL = "http://localhost:11434/api/chat"


def check_qwen_available() -> bool:
    """检查本地 qwen2.5vl 是否可用(默认关闭,防止 OOM)"""
    # qwen2.5vl OCR 内存消耗大,容易 OOM 导致进程被 kill。
    # 默认返回 False,优先使用 EasyOCR。如需启用 qwen,设置环境变量 ENABLE_QWEN=1。
    if os.environ.get("ENABLE_QWEN", "0") != "1":
        return False
    import urllib.request
    try:
        req = urllib.request.Request(
            "http://localhost:11434/api/tags",
            headers={"Content-Type": "application/json"}
        )
        with urllib.request.urlopen(req, timeout=5) as resp:
            data = json.loads(resp.read())
            models = [m["name"] for m in data.get("models", [])]
            if any("qwen" in m.lower() and "vl" in m.lower() for m in models):
                print(f"[INFO] qwen2.5vl 已启用(ENABLE_QWEN=1)", file=sys.stderr)
                return True
    except Exception:
        pass
    return False


QWEN_OCR_PROMPT = """你是一个专业的文档 OCR 提取助手。请仔细识别这张 PDF 页面的内容,并按以下 JSON 格式返回:
```json
{
  "page": <页码>,
  "blocks": [
    {"type": "text", "content": "<识别的文本段落>", "heading": "<如果是标题则填入标题文本,否则空字符串>"},
    {
      "type": "table",
      "content": "<表格的纯文本表示,每行用换行分隔>",
      "table_data": [
        {
          "cells": ["列1", "列2", "列3"],
          "col_span": [1, 1, 1],
          "row_span": [1, 1, 1],
          "cell_bbox": [[10, 20, 80, 35], [85, 20, 155, 35], [160, 20, 230, 35]]
        },
        {
          "cells": ["值1", "值2", "值3"],
          "col_span": [1, 2, 1],
          "row_span": [1, 1, 1],
          "cell_bbox": [[10, 40, 80, 55], [85, 40, 230, 55]]
        }
      ]
    }
  ]
}
```
重要规则:
1. 只输出 JSON,不要任何解释
2. table 块必须同时提供 content(纯文本)和 table_data(对象数组),缺一不可
3. table_data 每个单元格必须是字符串,空单元格用 "" 表示
4. cell_bbox(强烈建议提供):每个单元格的边界框坐标数组,单位为点数(pt),格式为 [left, bottom, right, top],坐标原点为页面左下角。
   - 若某行未提供 cell_bbox,则用前一行的 bbox 按行号递推估算(行间距默认20pt)
   - 若 cell_bbox 列数与 cells 列数不一致,以 cells 为准,bbox 多余部分忽略,少则按列均分
5. col_span / row_span 均为整数数组,长度与 cells 相同:
   - col_span[i]=N 表示该单元格从当前列起占据 N 列(跨列)
   - row_span[i]=N 表示该单元格从当前行起占据 N 行(跨行)
   - 跨行/跨列的单元格在其所跨越的后续位置不再重复出现
   - 解析规则:按行展开;若 col_span=[1,2,1],第1列正常,第2列跨2列(填入内容,后续位置跳过),第3列正常
6. 嵌套表头检测规则:
   - 若某行所有列的 row_span > 1,表示该行是上级表头行,其下各行属于该表头的子列
   - 若一行的多个连续列 col_span > 1 且内容简洁(≤4字符),这些列可能是下级表头
   - 典型嵌套结构:第1行["项目","  ","统计分析"] + 第2行["  ","A组","B组"]
7. 斜线表头检测规则:
   - 若表头某行存在单元格左下角(x1,y1)和右上角(x2,y2),且y1>y2,则该单元格为斜线表头
   - 斜线表头的 cells 数组内相邻非空元素的水平投影应有一定间隔(≥表格宽度的1/4)
8. 文本块如果包含标题(如"一、概述"),heading 字段填入标题文本
9. 识别所有中英文内容,数字、符号、化学式都要准确
10. 对识别质量进行评分,填入 confidence 字段(high/medium/low),评判依据:文字清晰度、表格结构完整性、是否有遗漏或模糊区域"""


def qwen_ocr_pdf(pdf_path: str, output_dir: str,
                lang: str = "zh,en",
                pages_to_skip: list = None) -> dict:
    """
    使用本地 qwen2.5vl 对扫描 PDF 进行 OCR
    逐页处理,每页独立获取VRAM → OCR → 立即释放VRAM

    pages_to_skip: 数字页列表(只跳过,不 OCR),节省时间
    """
    import urllib.request
    import base64
    import io

    _vm = None
    _result = {"success": False}

    # ---- VRAM 初始化(获取调度器引用,暂不占用)----
    try:
        sys.path.insert(0, "/home/wangyc/.openclaw/scripts")
        from vram_manager import VMgr
        _vm = VMgr()
    except Exception as e:
        print(f"[警告] VRAM 调度器初始化失败: {e},将直接执行 OCR...", file=sys.stderr)

    # ---- 正式 OCR 处理(外层 try 确保 VRAM 始终被恢复)----
    try:
        try:
            import pypdfium2 as pdfium
        except ImportError:
            print("[错误] pypdfium2 未安装,qwen2.5vl OCR 无法使用", file=sys.stderr)
            return _result

        try:
            doc = pdfium.PdfDocument(pdf_path)
            total_pages = len(doc)
            print(f"[INFO] PDF 共 {total_pages} 页,开始逐页 OCR(每页独立VRAM)...")
        except Exception as e:
            print(f"[错误] 无法读取 PDF: {e}", file=sys.stderr)
            return _result

        basename = Path(pdf_path).stem
        os.makedirs(output_dir, exist_ok=True)
        json_path = Path(output_dir) / f"{basename}.json"
        md_path = Path(output_dir) / f"{basename}.md"

        all_elements = []
        total_tables = 0
        total_paras = 0

        skip_set = set(pages_to_skip or [])
        for page_idx in range(total_pages):
            page_num = page_idx + 1
            if page_num in skip_set:
                continue  # 跳过数字页,只 OCR 扫描页

            # === VRAM 每页获取(渲染前让出,给ComfyUI机会)===
            _page_vram_acquired = False
            if _vm is not None:
                try:
                    if _vm.acquire_for_comfy(reason="opendataloader-qwen-ocr"):
                        _page_vram_acquired = True
                        print(f"[INFO] VRAM 已腾出(第{page_num}页),qwen2.5vl 可用")
                except Exception as _e:
                    print(f"[警告] VRAM 获取异常: {_e},继续尝试...", file=sys.stderr)

            try:
                page = doc.get_page(page_idx)
                pil_img = page.render(scale=72 / 72.0).to_pil()
                buf = io.BytesIO()
                pil_img.save(buf, format="JPEG", quality=75)
                img_b64 = base64.b64encode(buf.getvalue()).decode("ascii")
                pil_img.close()  # 立即释放PIL图像内存
            except Exception as _e:
                print(f"[警告] 第 {page_num} 页渲染失败: {_e}", file=sys.stderr)
                if _page_vram_acquired and _vm is not None:
                    try:
                        _vm.release_and_restore()
                    except Exception:
                        pass
                continue

            # 调用 Ollama(重试3次)
            content = None
            for attempt in range(3):
                try:
                    payload = {
                        "model": QWEN_MODEL,
                        "messages": [{"role": "user", "content": QWEN_OCR_PROMPT.format(
                            prompt_img=img_b64, lang=lang)}],
                        "stream": False,
                        "options": {"temperature": 0.01}
                    }
                    req = urllib.request.Request(
                        OLLAMA_URL,
                        data=json.dumps(payload).encode("utf-8"),
                        headers={"Content-Type": "application/json"},
                        method="POST"
                    )
                    with urllib.request.urlopen(req, timeout=180) as resp:
                        result_data = json.loads(resp.read())
                        content = result_data.get("message", {}).get("content", "")
                        ocr_done_reason = result_data.get("done_reason", "")
                        ocr_eval_count = result_data.get("eval_count", 0)
                        ocr_total_duration = result_data.get("total_duration", 0)
                        break
                except Exception as _e:
                    if attempt < 2:
                        import time as _time
                        _time.sleep(2)
                    else:
                        print(f"[警告] 第 {page_num} 页 OCR 失败: {_e}", file=sys.stderr)

            # === VRAM 每页立即释放(趁qwen2.5vl还在显存中)===
            if _page_vram_acquired and _vm is not None:
                try:
                    _vm.release_and_restore()
                except Exception as _e:
                    print(f"[警告] VRAM 恢复异常(页{page_num}): {_e}", file=sys.stderr)

            if not content:
                continue

            # 计算该页 OCR 置信度
            ocr_quality = "high"
            if ocr_done_reason and ocr_done_reason != "stop":
                ocr_quality = "medium"
            if ocr_eval_count > 0:
                avg_chars_per_token = len(content) / ocr_eval_count if ocr_eval_count else 0
                if avg_chars_per_token < 3 or avg_chars_per_token > 50:
                    ocr_quality = "low"
            page_confidence = ocr_quality

            # 解析 JSON
            try:
                if content.startswith('```'):
                    _lines_split = content.split('\n')
                    content = '\n'.join(_lines_split[1:] if _lines_split[0].startswith('```') else _lines_split)
                    if content.endswith('```'):
                        content = content[:-3].strip()
                start, end = content.find('{'), content.rfind('}') + 1
                if start >= 0 and end > start:
                    page_data = json.loads(content[start:end])
                else:
                    continue
            except Exception:
                continue

            for block in page_data.get("blocks", []):
                btype = block.get("type", "")
                bcontent = block.get("content", "").strip()
                if not bcontent:
                    continue
                if btype == "text":
                    heading = block.get("heading", "") or ""
                    total_paras += 1
                    elem = {"type": "paragraph", "content": bcontent,
                        "page_number": page_num, "paragraph_index": total_paras,
                        "ocr_confidence": page_confidence}
                    if heading:
                        elem["heading"] = heading
                    all_elements.append(elem)
                elif btype == "table":
                    total_tables += 1
                    elem = {"type": "table", "content": bcontent,
                        "page_number": page_num, "table_index": total_tables,
                        "ocr_confidence": page_confidence}
                    raw_td = block.get("table_data", [])
                    table_data, has_span = _parse_table_data(raw_td)
                    elem["table_data"] = table_data
                    elem["has_merged_cells"] = has_span
                    elem["_raw_td"] = raw_td
                    elem["_cell_bboxes"] = [row.get("cell_bbox") for row in raw_td if isinstance(row, dict)]
                    structure = _analyze_table_structure(raw_td)
                    elem["header_depth"] = structure.get("header_depth", 1)
                    elem["has_nested_headers"] = structure.get("has_nested_headers", False)
                    elem["has_diagonal_header"] = structure.get("has_diagonal_header", False)
                    elem["col_groups"] = structure.get("col_groups", [])
                    all_elements.append(elem)

            print(f"[INFO] 第 {page_num}/{total_pages} 页完成")

        # 保存结果
        output_json = {
            "doc_type": "scanned_pdf",
            "source": "qwen2.5vl-ocr",
            "total_pages": total_pages,
            "traceability": {
                "total_elements": len(all_elements),
                "total_paragraphs": total_paras,
                "total_tables": total_tables,
                "ocr_model": QWEN_MODEL,
            },
            "elements": all_elements,
            "flat_elements": all_elements,
        }
        with open(json_path, "w", encoding="utf-8") as f:
            json.dump(output_json, f, ensure_ascii=False, indent=2)

        md_lines = [f"# {basename}\n"]
        for elem in all_elements:
            pg = elem.get("page_number")
            md_lines.append(f"\n--- 第 {pg} 页 ---")
            md_lines.append(elem["content"])
        with open(md_path, "w", encoding="utf-8") as f:
            f.write('\n'.join(md_lines))

        print(f"[成功] qwen2.5vl OCR 完成: {len(all_elements)} 个元素")

        # 跨页表格拼接
        all_elements = _merge_consecutive_tables(all_elements)

        _result = {"success": True, "json_path": str(json_path), "md_path": str(md_path)}

    finally:
        # VRAM 最终恢复(清理残余状态)
        if _vm is not None:
            try:
                _vm.release_and_restore()
            except Exception:
                pass

    return _result


def convert_word_to_pdf(docx_path: str, output_dir: str = "/tmp") -> str:
    """
    用 LibreOffice 将 Word 文档转换为 PDF
    返回: PDF 文件的绝对路径
    """
    import tempfile

    pdf_dir = output_dir if output_dir else "/tmp"
    os.makedirs(pdf_dir, exist_ok=True)

    basename = Path(docx_path).stem + ".pdf"
    pdf_path = os.path.join(pdf_dir, basename)

    cmd = [
        "libreoffice", "--headless",
        "--convert-to", "pdf",
        "--outdir", pdf_dir,
        os.path.abspath(docx_path),
    ]

    print(f"[INFO] Word → PDF 转换中: {' '.join(cmd)}")
    result = subprocess.run(cmd, capture_output=True, text=True, timeout=120)

    if result.returncode != 0:
        raise RuntimeError(f"LibreOffice 转换失败: {result.stderr[:300]}")

    # LibreOffice 可能改名(如加编号),找最新生成的 pdf
    # 等待文件落盘
    time.sleep(1)

    # 精确路径:与 docx 同名的 pdf
    expected = os.path.join(pdf_dir, basename)
    if os.path.exists(expected):
        print(f"[INFO] PDF 生成成功: {expected}")
        return expected

    # 备选:找最新的 pdf 文件
    files = sorted(Path(pdf_dir).glob("*.pdf"), key=lambda p: p.stat().st_mtime, reverse=True)
    if files:
        print(f"[INFO] PDF 生成成功: {files[0]}")
        return str(files[0])

    raise RuntimeError("PDF 转换后未找到输出文件")


# ---------- PDF 提取结果增强(内容指纹) ----------
def add_content_fingerprint(data: dict) -> dict:
    """
    给 PDF 提取的 JSON 添加内容指纹,增强可追溯性。
    1. 扁平化 kids 嵌套结构 → flat_elements
    2. 添加 section_path(标题层级路径)
    3. 添加 table_index / paragraph_index(全局序号)
    4. 添加 content_preview(前100字符,用于快速定位)
    """
    flat = []
    heading_path = []  # 栈:当前标题路径

    def flatten(kids, flat_list, heading_path):
        for kid in kids:
            elem = kid.copy()
            kid_type = elem.get("type", "")

            # 维护标题路径
            if kid_type in ("heading", "header"):
                level = elem.get("heading level", 99)
                # Doctitle=0, Subtitle=1, 正文=2...
                # 简单映射:heading level 数值
                # 弹出更高级的标题
                while heading_path and heading_path[-1][0] >= level:
                    heading_path.pop()
                heading_path.append((level, elem.get("content", "")))
                elem["section_path"] = " > ".join(h for _, h in heading_path)
                flat_list.append(elem)
                # 递归处理 kids(子元素)
                if "kids" in elem:
                    flatten(elem.pop("kids"), flat_list, heading_path)
            else:
                flat_list.append(elem)
                if "kids" in elem:
                    flatten(elem.pop("kids"), flat_list, heading_path)

    if "kids" in data:
        flatten(data["kids"], flat, heading_path)

    # 全局序号
    table_count = 0
    para_count = 0
    img_count = 0

    for elem in flat:
        t = elem.get("type", "")
        content = elem.get("content", "")

        # 内容预览(前100字)
        elem["content_preview"] = content[:100] + ("..." if len(content) > 100 else "")

        if t == "table":
            table_count += 1
            elem["table_index"] = table_count
        elif t == "paragraph":
            para_count += 1
            elem["paragraph_index"] = para_count
        elif t in ("image", "figure"):
            img_count += 1
            elem["image_index"] = img_count

    # 元数据
    data["traceability"] = {
        "total_elements": len(flat),
        "total_tables": table_count,
        "total_paragraphs": para_count,
        "total_images": img_count,
        "fingerprint_version": "1.0",
    }
    data["flat_elements"] = flat
    return data


# ---------- 核心:合并 Word 内容 + PDF 位置信息 ----------
def normalize_text(text: str) -> str:
    """文本归一化:去空格、标点、换行,用于匹配"""
    if not text:
        return ""
    import unicodedata
    text = unicodedata.normalize("NFKC", text)
    text = re.sub(r"[\s\n\r\t\xa0·\-----]", "", text)
    text = text.lower()
    return text


def text_similarity(t1: str, t2: str) -> float:
    """两个文本的相似度(0~1)"""
    n1, n2 = normalize_text(t1), normalize_text(t2)
    if not n1 or not n2:
        return 0.0
    if n1 == n2:
        return 1.0
    # 简单:n1 在 n2 中或 n2 在 n1 中 → 高相似
    if n1 in n2 or n2 in n1:
        return 0.8
    # Jaccard
    s1, s2 = set(n1), set(n2)
    if not s1 or not s2:
        return 0.0
    return len(s1 & s2) / len(s1 | s2)


def extract_pdf_positions(pdf_flat: list) -> dict:
    """
    从 PDF 扁平元素列表中提取:
    - 标题层级路径(heading → section_path)
    - 页码 + bbox(heading / table / paragraph / image)
    返回结构:{(type, normalized_key): {page, bbox, section_path}}
    """
    pos_map = {}
    heading_path = []

    for elem in pdf_flat:
        t = elem.get("type", "")
        content = elem.get("content", "")

        if t == "heading":
            level = elem.get("heading level", 99)
            heading_path.append((level, content))
            heading_path.sort(key=lambda x: x[0])
            path_str = " > ".join(h for _, h in heading_path)
            key = normalize_text(content)
            pos_map[("heading", key)] = {
                "page": elem.get("page number"),
                "bbox": elem.get("bounding box"),
                "section_path": path_str,
            }
        elif t in ("table", "paragraph", "image", "figure"):
            key = normalize_text(content)[:50]
            if ("table", key) not in pos_map:
                pos_map[(t, key)] = {
                    "page": elem.get("page number"),
                    "bbox": elem.get("bounding box"),
                }

    return pos_map


def find_best_pdf_match(docx_elem: dict, pdf_pos_map: dict,
                        seen_tables: set, seen_paras: set) -> dict:
    """
    为单个 docx 元素在 pdf_flat 中找最佳匹配的 position 信息。
    匹配策略:
      - 标题:按 text 精确匹配(normalize 后)
      - 表格:按内容相似度匹配(取第一行文字)
      - 段落:按内容相似度匹配(跳过已匹配的)
    返回:{page, bbox, section_path} 或 {}
    """
    t = docx_elem.get("type", "")
    content = docx_elem.get("content", "")

    norm_key = normalize_text(content)

    if t == "heading":
        for key, info in pdf_pos_map.items():
            if key[0] == "heading" and norm_key and key[1] == norm_key:
                return {"page": info["page"], "bbox": info["bbox"],
                        "section_path": info["section_path"]}
        # 模糊匹配
        for key, info in pdf_pos_map.items():
            if key[0] == "heading" and norm_key and text_similarity(content, key[1]) > 0.7:
                return {"page": info["page"], "bbox": info["bbox"],
                        "section_path": info["section_path"]}

    elif t == "table":
        # 取表格第一行做匹配键
        lines = content.strip().split("\n")
        if len(lines) >= 2:
            header = lines[1] if lines[0].startswith("|") else lines[0]
            norm_header = normalize_text(header)[:50]
            for key, info in pdf_pos_map.items():
                if key[0] == "table" and norm_header and key[1] == norm_header:
                    if key not in seen_tables:
                        seen_tables.add(key)
                        return {"page": info["page"], "bbox": info["bbox"]}
        # 模糊匹配
        for key, info in pdf_pos_map.items():
            if key[0] == "table" and key not in seen_tables:
                seen_tables.add(key)
                return {"page": info["page"], "bbox": info["bbox"]}

    elif t == "paragraph":
        if not norm_key:
            return {}
        for key, info in pdf_pos_map.items():
            if key[0] == "paragraph" and key not in seen_paras:
                sim = text_similarity(content, key[1])
                if sim > 0.7:
                    seen_paras.add(key)
                    return {"page": info["page"], "bbox": info["bbox"]}

    return {}


def merge_docx_and_pdf(docx_data: dict, pdf_flat: list) -> dict:
    """
    核心合并函数:以 docx 内容为主体,为每个元素附加 PDF 位置信息。
    相同信息以更可靠来源为准:表格/文本内容以 docx 为准,页码/bbox 以 PDF 为准。
    """
    # 构建 PDF position 字典
    pdf_pos_map = extract_pdf_positions(pdf_flat)
    seen_tables = set()
    seen_paras = set()

    # 全局序号
    table_count = 0
    para_count = 0
    img_count = 0

    # 收集 section_path(从 docx 的标题层级)
    heading_path = []

    merged_elements = []
    for elem in docx_data.get("elements", []):
        elem = dict(elem)  # 浅拷贝
        t = elem.get("type", "")

        # 更新标题路径
        if t == "heading":
            level = elem.get("heading_level", 99)
            content = elem.get("content", "")
            while heading_path and heading_path[-1][0] >= level:
                heading_path.pop()
            heading_path.append((level, content))
            elem["section_path"] = " > ".join(h for _, h in heading_path)

        # 内容预览
        elem["content_preview"] = elem.get("content", "")[:100]
        if len(elem.get("content", "")) > 100:
            elem["content_preview"] += "..."

        # 在 PDF 中找位置(最佳匹配)
        pos = find_best_pdf_match(elem, pdf_pos_map, seen_tables, seen_paras)
        if pos:
            elem["page"] = pos.get("page")
            elem["bbox"] = pos.get("bbox")

        # 序号
        if t == "table":
            table_count += 1
            elem["table_index"] = table_count
        elif t == "paragraph":
            para_count += 1
            elem["paragraph_index"] = para_count

        merged_elements.append(elem)

    # 元数据
    result = {
        "doc_type": "word",
        "source": "word-to-pdf-merged",
        "total_elements": len(merged_elements),
        "total_tables": table_count,
        "total_paragraphs": para_count,
        "elements": merged_elements,
        # PDF 位置信息(备用)
        "_pdf_position_hints": {
            "total_from_pdf": len(pdf_flat),
            "tables_matched": len(seen_tables),
            "paras_matched": len(seen_paras),
        }
    }
    return result


def convert_word_to_markdown(docx_path: str) -> str:
    """生成可读 Markdown(基于 docx,与合并后 JSON 对应)"""
    from docx import Document
    doc = Document(docx_path)
    lines = []
    heading_path = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style_name = para.style.name if para.style else ""
        if style_name.startswith("Heading") or (para.style and "Heading" in para.style.name):
            level = 1
            for c in style_name:
                if c.isdigit():
                    level = int(c)
                    break
            lines.append(f"{'#' * level} {text}")
            # 维护路径
            while heading_path and heading_path[-1][0] >= level:
                heading_path.pop()
            heading_path.append((level, text))
        else:
            lines.append(text)

    for table in doc.tables:
        rows_text = []
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            rows_text.append(row_text)
        if rows_text:
            lines.append("")
            lines.append(rows_text[0])
            lines.append(" | ".join(["---"] * len(rows_text[0].split(" | "))))
            for row in rows_text[1:]:
                lines.append(row)

    return "\n".join(lines)


# ---------- Excel 提取 ----------
def _make_serializable(val):
    """将 Excel 单元格值转为 JSON 可序列化类型"""
    import datetime
    if val is None:
        return ""
    if isinstance(val, (datetime.datetime, datetime.date, datetime.time)):
        return val.isoformat()
    if isinstance(val, float):
        # 截断浮点数精度
        return round(val, 6)
    return val


def extract_excel_to_json(xlsx_path: str) -> dict:
    """
    用 openpyxl 提取 Excel 完整内容:
    - 工作表列表 + 表头
    - 每个 sheet 的数据区域(转为 table 结构)
    - 命名区域 / 表格对象
    - chart 图表(提取图表类型 + 关联数据表)
    输出格式与 PDF/Word 统一(elements 数组)。
    """
    import openpyxl
    from openpyxl.utils import get_column_letter
    from openpyxl.chart import BarChart, LineChart, PieChart, ScatterChart, Reference

    wb = openpyxl.load_workbook(xlsx_path, data_only=True)
    elements = []
    sheet_count = 0
    table_count = 0
    chart_count = 0

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        sheet_count += 1

        # 跳过空 sheet
        if ws.max_row < 2 or ws.max_column < 2:
            continue

        # 判断 sheet 是否为表格(是否有数据的起始行)
        # 使用 span-aware 提取(支持 colspan/rowspan 合并单元格)
        rows_data = _extract_excel_table_with_spans(ws)

        if not rows_data:
            continue

        # 检测表头(第一行)
        headers = rows_data[0]["cells"] if rows_data else []

        # 检测是否有合并单元格
        has_merged = any(
            any(cs > 1 or rs > 1
                for cs, rs in zip(r.get("col_span", []), r.get("row_span", [])))
            for r in rows_data
        )
        has_nested_headers = False
        if len(rows_data) >= 2:
            first_col_spans = rows_data[0].get("col_span", [])
            if any(cs > 1 for cs in first_col_spans):
                has_nested_headers = True

        # 构建 Markdown 表格(colspan 横向展开)
        table_content = _table_to_markdown_with_spans(rows_data)

        # 保留纯数据行(供 AI 直接分析用)
        serializable_rows = [r["cells"] for r in rows_data]

        elements.append({
            "type": "table",
            "location": {
                "excel": sheet_name,           # 主要标识:原Excel Sheet名
                "pdf_page": None,               # PDF页码(匹配到则填入,否则None)
            },
            "table_index": table_count + 1,
            "content": table_content,
            "content_preview": table_content[:100],
            "row_count": len(rows_data),
            "col_count": max(len(r["cells"]) for r in rows_data) if rows_data else 0,
            "headers": headers,
            "data_rows": serializable_rows,     # 结构化数组,AI 可直接分析
            # 统一 table_data 格式(含 colspan/row_span,与 Word/qwen2.5vl 一致)
            "table_data": rows_data,
            "has_merged_cells": has_merged,
            "has_nested_headers": has_nested_headers,
            "section_path": f"Excel > {Path(xlsx_path).stem} > {sheet_name}",
        })
        table_count += 1

        # 提取 chart 信息
        if hasattr(ws, "_charts"):
            for chart in ws._charts:
                chart_count += 1
                chart_type = type(chart).__name__
                chart_title = chart.title if hasattr(chart, "title") and chart.title else ""

                # 提取图表关联的数据范围
                data_desc = ""
                if hasattr(chart, "data"):
                    refs = chart.data
                    if refs:
                        data_desc = f"数据范围: {refs}"

                # 构建图表的摘要数据(将图表还原为简化的数据表)
                sample_data = []
                if hasattr(chart, "series") and chart.series:
                    for series in chart.series:
                        sname = series.title if hasattr(series, "title") and series.title else f"系列{chart_count}"
                        sample_data.append(f"系列: {sname}")

                elements.append({
                    "type": "chart",
                    "location": {
                        "excel": sheet_name,           # 主要标识:原Excel Sheet名
                        "pdf_page": None,               # PDF页码(匹配到则填入,否则None)
                    },
                    "chart_index": chart_count,
                    "chart_type": chart_type,
                    "chart_title": str(chart_title),
                    "content": f"[{chart_type}] {chart_title} | {data_desc}",
                    "content_preview": f"{chart_type}: {chart_title}",
                    "section_path": f"Excel > {Path(xlsx_path).stem} > {sheet_name} > 图表",
                })

    # 命名区域
    defined_names = []
    if wb.defined_names:
        for name, defn in wb.defined_names.items():
            defined_names.append({
                "name": name,
                "value": _make_serializable(getattr(defn, "value", "")),
                "attr_text": str(getattr(defn, "attr_text", "")),
            })

    return {
        "doc_type": "excel",
        "source": "openpyxl",
        "total_sheets": sheet_count,
        "total_tables": table_count,
        "total_charts": chart_count,
        "defined_names": defined_names,
        "sheets": wb.sheetnames,
        "elements": elements,
    }


def extract_excel_to_markdown(xlsx_path: str, page_hints: dict = None) -> str:
    """生成 Excel 的可读 Markdown(每个 sheet 转为 Markdown 表格)"""
    import openpyxl

    wb = openpyxl.load_workbook(xlsx_path, data_only=True)
    page_hints = page_hints or {}
    lines = [f"# {Path(xlsx_path).stem}\n"]

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        page_hint = page_hints.get(sheet_name, "")
        pg_note = f" _(PDF 第{page_hint}页)_" if page_hint else ""
        lines.append(f"\n## Sheet: {sheet_name}{pg_note}\n")

        rows_data = _extract_excel_table_with_spans(ws)
        if not rows_data:
            lines.append("_(空表)_\n")
            continue

        lines.append(_table_to_markdown_with_spans(rows_data))
        lines.append("")

    return "\n".join(lines)


# ---------- Word 提取 ----------
    """用 python-docx 提取 Word 内容,输出与 PDF 相同的 JSON 结构"""
    from docx import Document

    doc = Document(docx_path)
    elements = []
    page = 1

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # 判断段落类型
        style_name = para.style.name.lower() if para.style else ""

        if style_name.startswith("heading") or para.style and "Heading" in para.style.name:
            # 提取标题级别
            level = 1
            for c in para.style.name:
                if c.isdigit():
                    level = int(c)
                    break
            elements.append({
                "type": "heading",
                "content": text,
                "heading_level": level,
                "page": page,
                "bbox": None,
            })
        else:
            elements.append({
                "type": "paragraph",
                "content": text,
                "page": page,
                "bbox": None,
            })

    # 提取表格
    for table in doc.tables:
        rows_data = []
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells]
            rows_data.append(row_texts)

        if rows_data:
            # 构建 Markdown 表格
            md_table = ""
            if rows_data:
                # 表头
                md_table += "| " + " | ".join(rows_data[0]) + " |\n"
                # 分隔线
                md_table += "| " + " | ".join(["---"] * len(rows_data[0])) + " |\n"
                # 数据行
                for row in rows_data[1:]:
                    md_table += "| " + " | ".join(row) + " |\n"

            elements.append({
                "type": "table",
                "content": md_table.strip(),
                "page": page,
                "bbox": None,
            })

    # 转为与 PDF 输出兼容的 JSON
    return {
        "doc_type": "word",
        "total_pages": 1,
        "elements": elements,
    }


def extract_word_to_markdown(docx_path: str) -> str:
    """用 python-docx 提取 Word 内容,输出 Markdown(与 PDF 提取的 .md 格式一致)"""
    from docx import Document

    doc = Document(docx_path)
    lines = []
    in_table = False

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            if in_table:
                lines.append("")
                in_table = False
            continue

        style_name = para.style.name if para.style else ""

        if style_name.startswith("Heading") or (para.style and "Heading" in para.style.name):
            # 标题
            level = 1
            for c in style_name:
                if c.isdigit():
                    level = int(c)
                    break
            lines.append(f"{'#' * level} {text}")
            in_table = False
        else:
            lines.append(text)
            in_table = False

    # 处理表格(支持 colspan/row_span 渲染)
    for table in doc.tables:
        rows_data = _extract_docx_table_with_spans(table)
        if rows_data:
            lines.append("")
            lines.append(_table_to_markdown_with_spans(rows_data))

    return "\n".join(lines)


def extract_word_to_json(docx_path: str,
                            extract_page_numbers: bool = True) -> dict:
    """
    用 python-docx 提取 Word 内容,输出 JSON 格式(与 PDF 提取的 JSON 格式一致)。

    参数:
      extract_page_numbers: 是否提取真实页码(内部自动 Word→PDF→opendataloader→页码匹配)。
                              关闭时 page=1(纯 docx 提取,无外部依赖)。
    """
    from docx import Document

    doc = Document(docx_path)
    elements = []
    paragraph_index = 0
    table_index = 0

    # ---- 提取段落(暂时不填 page) ----
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style_name = para.style.name if para.style else ""
        heading_level = None
        if style_name.startswith("Heading") or (para.style and "Heading" in para.style.name):
            for c in style_name:
                if c.isdigit():
                    heading_level = int(c)
                    break
            if heading_level is None:
                heading_level = 1
        elements.append({
            "type": "heading" if heading_level else "paragraph",
            "content": text,
            "page": None,       # 稍后填充
            "heading_level": heading_level,
            "paragraph_index": paragraph_index,
            "table_index": None,
            "section_path": "",
            "content_preview": text[:100] if len(text) > 100 else text,
            "original_word": docx_path,
        })
        paragraph_index += 1

    # ---- 提取表格(带 colspan/row_span,暂时不填 page) ----
    for table in doc.tables:
        rows_data = _extract_docx_table_with_spans(table)
        has_merged = any(
            any(cs > 1 or rs > 1
                for cs, rs in zip(r.get("col_span", []), r.get("row_span", [])))
            for r in rows_data
        )
        has_nested_headers = False
        if len(rows_data) >= 2:
            first_col_spans = rows_data[0].get("col_span", [])
            if any(cs > 1 for cs in first_col_spans):
                has_nested_headers = True
        if rows_data:
            md_table = _table_to_markdown_with_spans(rows_data)
            elements.append({
                "type": "table",
                "content": md_table,
                "page": None,
                "heading_level": None,
                "paragraph_index": None,
                "table_index": table_index,
                "section_path": "",
                "content_preview": f"表格 {table_index + 1}",
                "original_word": docx_path,
                "table_data": rows_data,
                "has_merged_cells": has_merged,
                "has_nested_headers": has_nested_headers,
            })
            table_index += 1

    # ---- Word → PDF → opendataloader → 页码匹配 ----
    page_map = {}          # element_index -> real_page
    section_paths = {}     # element_index -> section_path_str
    detected_pages = 1

    if extract_page_numbers:
        try:
            # Step 1: Word → PDF
            pdf_path = convert_word_to_pdf(docx_path, output_dir="/tmp")
            # Step 2: opendataloader 提取 PDF 页码位置
            env = os.environ.copy()
            env["JAVA_HOME"] = JAVA_HOME
            env["PATH"] = JAVA_HOME + "/bin:" + env.get("PATH", "")
            cmd = [
                sys.executable, "-m", "opendataloader_pdf",
                pdf_path, "-o", "/tmp", "-f", "json", "-q"
            ]
            subprocess.run(cmd, capture_output=True, text=True, env=env, timeout=300)
            import glob as _glob
            pdf_jsons = _glob.glob("/tmp/*.json")
            pdf_flat = []
            for jf in pdf_jsons:
                try:
                    with open(jf, encoding="utf-8") as f:
                        d = json.load(f)
                    kids_raw = d.get("kids", [])
                    if not kids_raw:
                        def _flat(k, out):
                            for item in k:
                                out.append(item)
                                if "kids" in item:
                                    _flat(item["kids"], out)
                        _flat(d.get("kids", []), kids_raw)
                    pdf_flat.extend(kids_raw)
                except Exception:
                    pass
            # Step 3: 建立 PDF position map
            pdf_pos_map = extract_pdf_positions(pdf_flat)
            detected_pages = max(
                [e.get("page number", 1) for e in pdf_flat] or [1])
            # Step 4: 逐元素匹配页码
            seen_tables_set = set()
            seen_paras_set = set()
            for i, elem in enumerate(elements):
                matched = find_best_pdf_match(elem, pdf_pos_map,
                                              seen_tables_set, seen_paras_set)
                if matched:
                    pg = matched.get("page")
                    if pg:
                        page_map[i] = pg
                    sp = matched.get("section_path")
                    if sp:
                        section_paths[i] = sp
            # 清理临时 PDF
            Path(pdf_path).unlink(missing_ok=True)
            for jf in pdf_jsons:
                Path(jf).unlink(missing_ok=True)
        except Exception as e:
            print(f"[警告] Word 页码提取失败: {e}, page=1", file=sys.stderr)

    # ---- 回填 page 和 section_path ----
    for i, elem in enumerate(elements):
        elem["page"] = page_map.get(i, 1)
        if section_paths.get(i):
            elem["section_path"] = section_paths[i]

    # ---- kids:直接从 elements 复制(含 page_number) ----
    kids = []
    for e in elements:
        ke = {
            "type": e["type"],
            "page number": e.get("page", 1),
            "content": e.get("content", ""),
            "heading_level": e.get("heading_level"),
            "bounding box": [0, 0, 0, 0],
        }
        if e["type"] == "table":
            ke["table_data"] = e.get("table_data", [])
        kids.append(ke)

    return {
        "doc_type": "word",
        "original_word": docx_path,
        "number of pages": detected_pages,
        "kids": kids,
        "traceability": {
            "total_elements": len(elements),
            "total_paragraphs": paragraph_index,
            "total_tables": table_index,
        },
        "flat_elements": elements
    }


def convert_word(input_path: str, output_dir: str, output_format: str = "markdown,json") -> dict:
    """转换 Word 文档,统一输出格式"""
    os.makedirs(output_dir, exist_ok=True)

    basename = Path(input_path).stem
    formats = output_format.split(",")

    result = {
        "success": True,
        "input_path": input_path,
        "output_dir": output_dir,
        "files_created": [],
        "mode": "word-direct",
    }

    if "json" in formats:
        json_data = extract_word_to_json(input_path)
        json_path = os.path.join(output_dir, f"{basename}.json")
        with open(json_path, "w", encoding="utf-8") as f:
            json.dump(json_data, f, ensure_ascii=False, indent=2)
        result["files_created"].append(json_path)

    if "markdown" in formats:
        md_data = extract_word_to_markdown(input_path)
        md_path = os.path.join(output_dir, f"{basename}.md")
        with open(md_path, "w", encoding="utf-8") as f:
            f.write(f"# {basename}\n\n")
            f.write(md_data)
        result["files_created"].append(md_path)

    if "text" in formats:
        from docx import Document
        doc = Document(input_path)
        text_path = os.path.join(output_dir, f"{basename}.txt")
        with open(text_path, "w", encoding="utf-8") as f:
            for para in doc.paragraphs:
                f.write(para.text + "\n")
        result["files_created"].append(text_path)

    # 如果格式都不对,默认输出 markdown 和 json
    if not result["files_created"]:
        json_data = extract_word_to_json(input_path)
        md_data = extract_word_to_markdown(input_path)

        json_path = os.path.join(output_dir, f"{basename}.json")
        with open(json_path, "w", encoding="utf-8") as f:
            json.dump(json_data, f, ensure_ascii=False, indent=2)
        result["files_created"].append(json_path)

        md_path = os.path.join(output_dir, f"{basename}.md")
        with open(md_path, "w", encoding="utf-8") as f:
            f.write(f"# {basename}\n\n")
            f.write(md_data)
        result["files_created"].append(md_path)

    return result


def find_java():
    """检查 Java 是否可用"""
    try:
        result = subprocess.run(
            [JAVA_BIN, "-version"],
            capture_output=True, text=True, timeout=5
        )
        return True
    except Exception:
        return False


def classify_pages(pdf_path: str) -> dict:
    """
    逐页分类 PDF,每页独立判断是 digital 还是 scanned
    返回: {
        "total_pages": int,
        "digital_pages": [int, ...],   # 有文字的页(1-indexed)
        "scanned_pages": [int, ...],     # 无文字的页(1-indexed)
        "per_page": {page_num: "digital"|"scanned"},  # 每页单独结果
        "lang": "auto"|"zh"|"en",
    }
    """
    result = {
        "total_pages": 0,
        "digital_pages": [],
        "scanned_pages": [],
        "per_page": {},
        "lang": "auto",
    }

    try:
        from pypdf import PdfReader
        reader = PdfReader(pdf_path)
        total = len(reader.pages)
        result["total_pages"] = total

        all_text = ""
        for i, page in enumerate(reader.pages):
            page_num = i + 1
            text = page.extract_text() or ""
            chars = len(text.strip())
            if chars > 50:
                result["per_page"][page_num] = "digital"
                result["digital_pages"].append(page_num)
                all_text += text
            else:
                result["per_page"][page_num] = "scanned"
                result["scanned_pages"].append(page_num)

        # 语言检测
        chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', all_text))
        english_chars = len(re.findall(r'[a-zA-Z]', all_text))
        if chinese_chars > english_chars * 0.3 and chinese_chars > 20:
            result["lang"] = "zh"
        elif english_chars > 50:
            result["lang"] = "en"

    except ImportError:
        print("[警告] pypdf 未安装,逐页分类失败", file=sys.stderr)
    except Exception as e:
        print(f"[警告] 逐页分类异常: {e}", file=sys.stderr)

    return result


def detect_pdf_type(pdf_path: str) -> dict:
    """
    检测 PDF 类型(整体判断,用于兼容模式)
    返回: {
        "type": "digital" | "scanned" | "mixed",
        "lang": "auto" | "zh" | "en" | "mixed",
        "has_complex_tables": bool,
        "has_formulas": bool,
        "has_charts": bool,
        "total_pages": int,
        "text_pages": int,
        "scanned_pages": int,
    }
    """
    result = {
        "type": "digital",
        "lang": "auto",
        "has_complex_tables": False,
        "has_formulas": False,
        "has_charts": False,
        "total_pages": 0,
        "text_pages": 0,
        "scanned_pages": 0,
    }

    try:
        from pypdf import PdfReader
        reader = PdfReader(pdf_path)
        result["total_pages"] = len(reader.pages)

        # 检查前5页(统计有文字的页数)
        # 采样策略:取文档前中后三个区段,避免封面/目录页干扰
        total_pages = len(reader.pages)
        sample_indices = (
            list(range(0, min(3, total_pages))) +       # 前段:第1-3页
            list(range(total_pages // 2 - 1, total_pages // 2 + 1)) +  # 中段
            list(range(max(0, total_pages - 3), total_pages))          # 后段:最后3页
        )
        sample_indices = sorted(set(i for i in sample_indices if i < total_pages))

        text_count = 0
        scanned_count = 0

        for i in sample_indices:
            page = reader.pages[i]
            text = page.extract_text() or ""
            chars = len(text.strip())
            if chars > 50:
                text_count += 1
            else:
                scanned_count += 1

        sample_size = len(sample_indices)
        scanned_pages = scanned_count
        result["scanned_pages"] = scanned_pages
        result["text_pages"] = text_count

        # 判断类型(任一采样页为扫描版即触发 OCR)
        if scanned_pages == 0:
            result["type"] = "digital"
        elif text_count == 0:
            result["type"] = "scanned"
        else:
            result["type"] = "mixed"  # 混合:部分扫描+部分文字

        # 语言检测(基于采样页提取的文字)
        all_text = ""
        for i in sample_indices:
            all_text += (reader.pages[i].extract_text() or "")

        # 简单的中英文检测
        chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', all_text))
        english_chars = len(re.findall(r'[a-zA-Z]', all_text))

        if chinese_chars > english_chars * 0.3 and chinese_chars > 20:
            result["lang"] = "zh"
        elif english_chars > 50:
            result["lang"] = "en"
        else:
            result["lang"] = "auto"

    except ImportError:
        print("[警告] pypdf 未安装,无法自动检测,使用默认配置(digital)", file=sys.stderr)
    except Exception as e:
        print(f"[警告] PDF 检测失败: {e},使用默认配置(digital)", file=sys.stderr)

    return result


def merge_per_page_results(digital_result: dict, ocr_result: dict) -> dict:
    """
    合并两个来源的结果:
    - digital_result: opendataloader 处理数字页的输出(kids 结构)
    - ocr_result: qwen2.5vl 处理扫描页的输出(kids 结构)
    按 page_number 排序,合并为统一的 kids 结构
    """
    digital_kids = digital_result.get("kids", []) if digital_result else []
    ocr_kids = ocr_result.get("kids", []) if ocr_result else []

    # 全部收集,按 page_number 分组
    all_elements = {}

    for elem in digital_kids:
        pg = elem.get("page_number", 1)
        if pg not in all_elements:
            all_elements[pg] = []
        all_elements[pg].append(elem)

    for elem in ocr_kids:
        pg = elem.get("page_number", 1)
        if pg not in all_elements:
            all_elements[pg] = []
        all_elements[pg].append(elem)

    # 按页号排序,收集所有元素
    merged_kids = []
    total_tables = 0
    total_paras = 0
    for pg in sorted(all_elements.keys()):
        for elem in all_elements[pg]:
            merged_kids.append(elem)
            if elem.get("type") == "table":
                total_tables += 1
                if elem.get("has_merged_cells"):
                    total_merged_tables += 1
                if elem.get("has_nested_headers"):
                    total_nested_tables += 1
                if elem.get("has_diagonal_header"):
                    total_diagonal_tables += 1
            elif elem.get("type") == "paragraph":
                total_paras += 1

    # 构建输出
    output = {
        "doc_type": "pdf",
        "source": "per-page-adaptive",
        "total_pages": len(set(
            e.get("page_number", 1) for e in merged_kids
        )),
        "total_tables": total_tables,
        "total_paragraphs": total_paras,
        "kids": merged_kids,
        "total_merged_tables": total_merged_tables,
        "total_nested_tables": total_nested_tables,
        "total_diagonal_tables": total_diagonal_tables,
        "_digital_kids_count": len(digital_kids),
        "_ocr_kids_count": len(ocr_kids),
    }

    # 合并检测元数据
    if digital_result:
        for key in ["lang", "has_formulas", "has_charts"]:
            if key in digital_result:
                output[key] = digital_result[key]
    if ocr_result and "lang" not in output:
        output["lang"] = ocr_result.get("lang", "auto")

    return output


def process_pdf_per_page(pdf_path: str, output_dir: str,
                         output_format: str = "markdown,json",
                         force_qwen: bool = False,
                         skip_server: bool = False) -> dict:
    """
    逐页独立判断 + 最佳方法处理：
    对每一页单独判断类型（digital/scanned），并选择最适合的提取方法。

    处理逻辑（每页独立决策）：
      - scanned 页面（无文字） → qwen2.5vl OCR
      - digital 页面有表格（含疑似合并单元格） → qwen2.5vl OCR（精确结构）
      - digital 页面纯文本 → opendataloader fast（快）
      - 提取失败 → 自动降级

    执行策略（高效版，非逐页重复调用）：
      1. 逐页分类（pypdf，无成本）
      2. digital 页批量 opendataloader 一次扫描 → 找出有表格的页
      3. 有表格的 digital 页 + 全部 scanned 页 → qwen2.5vl（VRAM 感知）
      4. 纯文本 digital 页 → opendataloader fast 结果直接用
    """
    result = {
        "success": False,
        "pdf_path": pdf_path,
        "output_dir": output_dir,
        "mode_used": "per-page-adaptive",
    }

    # ---- Step 1: 逐页分类（pypdf，无成本） ----
    print(f"[Step 1/4] 逐页分类 PDF 类型...")
    classification = classify_pages(pdf_path)
    digital_pages = classification["digital_pages"]
    scanned_pages = classification["scanned_pages"]
    lang = classification.get("lang", "auto")
    per_page = classification.get("per_page", {})
    total = classification["total_pages"]
    print(f"[      ] 总页数: {total} | digital: {len(digital_pages)} | scanned: {len(scanned_pages)}")

    if not digital_pages and not scanned_pages:
        print("[错误] 无法读取 PDF 页", file=sys.stderr)
        return result

    # ---- Step 2: digital 页批量扫描（一次性找出有表格的页） ----
    page_elements = {}   # {page_num: [elem, ...]}
    page_has_table = {}   # {page_num: True/False}
    page_text_kids = {}  # opendataloader 的 digital 页文本结果

    if digital_pages:
        print(f"[Step 2/4] digital 页批量快速扫描（{len(digital_pages)} 页）...")
        env = os.environ.copy()
        env["JAVA_HOME"] = JAVA_HOME
        env["PATH"] = JAVA_HOME + "/bin:" + env.get("PATH", "")
        cmd = build_opendataloader_cmd(pdf_path, output_dir,
                                       {"type": "digital", "lang": lang, "total_pages": total},
                                       output_format)
        try:
            proc = subprocess.run(cmd, capture_output=True, text=True, env=env, timeout=300)
            if proc.returncode == 0:
                import glob as _glob
                basename_str = Path(pdf_path).stem
                json_files = _glob.glob(os.path.join(output_dir, f"{basename_str}*.json"))
                for jf in json_files:
                    with open(jf, encoding="utf-8") as f:
                        od_data = json.load(f)
                    kids_raw = od_data.get("kids", [])
                    if not kids_raw:
                        def _flat(kids, out):
                            for k in kids:
                                out.append(k)
                                if "kids" in k:
                                    _flat(k["kids"], out)
                        _flat(od_data.get("kids", []), kids_raw)
                    for k in kids_raw:
                        pg = k.get("page number", 1)
                        if pg not in page_elements:
                            page_elements[pg] = []
                        page_elements[pg].append(k)
                        if k.get("type") == "table":
                            # 检查列数是否均匀（不均匀→疑似合并单元格）
                            td = k.get("table_data")
                            if td and isinstance(td, list) and len(td) > 1:
                                col_counts = [len(r) for r in td if isinstance(r, list)]
                                if col_counts and (min(col_counts) < max(col_counts)):
                                    page_has_table[pg] = True
                                    continue
                            page_has_table[pg] = True
        except Exception as e:
            print(f"[警告] digital 页批量扫描失败: {e}，所有 digital 页当纯文本处理", file=sys.stderr)

        # 分类 digital 页：有表格 vs 纯文本
        digital_with_tables = [pg for pg in digital_pages if page_has_table.get(pg)]
        digital_text_only   = [pg for pg in digital_pages if pg not in page_has_table]
        print(f"[      ] digital 扫描结果：{len(digital_with_tables)} 页有表格，{len(digital_text_only)} 页纯文本")
    else:
        digital_with_tables = []
        digital_text_only   = []

    # ---- Step 3: qwen2.5vl 处理需要 OCR 的页（有表格的 digital + 全部 scanned） ----
    ocr_pages = list(set(digital_with_tables + scanned_pages))
    ocr_kids_by_page = {}  # {page_num: [elem, ...]}

    if ocr_pages:
        print(f"[Step 3/4] {len(ocr_pages)} 页需要 qwen2.5vl OCR（VRAM 感知队列）...")
        _got_vram = False
        _vm = None
        try:
            sys.path.insert(0, str(Path("/home/wangyc/.openclaw/scripts")))
            from vram_manager import VMgr
            _vm = VMgr()
            import time as _time
            _waited = 0
            while _waited < 300:
                _status = _vm.status()
                _vram_free = _status.get("vram_free_mb", 0)
                _state = _status.get("state", "?")
                if _vram_free >= 18000:
                    if _vm.acquire_for_comfy(reason="opendataloader-qwen-ocr"):
                        _got_vram = True
                        break
                if _waited == 0:
                    print(f"[      ] VRAM 排队等待（free={_vram_free}MB）...")
                _time.sleep(10)
                _waited += 10
        except Exception as _e:
            print(f"[警告] VRAM 队列异常: {_e}", file=sys.stderr)
        finally:
            if _vm is not None:
                try:
                    _vm.release_and_restore()
                except Exception:
                    pass

        # 逐页渲染 + qwen2.5vl OCR（每页独立，精准控制）
        for pg_num in sorted(ocr_pages):
            print(f"[      ]   处理第 {pg_num}/{total} 页 → qwen2.5vl")
            render_dir = os.path.join(output_dir, f"_page_{pg_num:04d}_render")
            os.makedirs(render_dir, exist_ok=True)
            img_map = _render_pdf_pages_as_images(pdf_path, [pg_num], render_dir)
            if not img_map:
                print(f"[警告]   第 {pg_num} 页渲染失败，跳过")
                continue
            img_path = img_map.get(pg_num)
            if not img_path:
                continue
            elems = _ocr_images_with_qwen([img_path], lang=lang)
            for e in elems:
                e["page_number"] = pg_num
            ocr_kids_by_page[pg_num] = elems

    # ---- Step 4: 合并所有结果 ----
    merged_kids = []
    total_tables = total_paras = 0
    total_merged = total_nested = total_diagonal = 0

    def _count_elem(elem):
        nonlocal total_tables, total_paras, total_merged, total_nested, total_diagonal
        if elem.get("type") == "table":
            total_tables += 1
            if elem.get("has_merged_cells"):
                total_merged += 1
            if elem.get("has_nested_headers"):
                total_nested += 1
            if elem.get("has_diagonal_header"):
                total_diagonal += 1
        elif elem.get("type") in ("paragraph", "heading"):
            total_paras += 1

    # digital 纯文本页 → opendataloader 结果
    for pg in sorted(digital_text_only):
        for elem in page_elements.get(pg, []):
            merged_kids.append(elem)
            _count_elem(elem)

    # digital 有表格页 → qwen2.5vl 结果（覆盖 opendataloader）
    for pg in sorted(digital_with_tables):
        if pg in ocr_kids_by_page:
            for elem in ocr_kids_by_page[pg]:
                merged_kids.append(elem)
                _count_elem(elem)
        else:
            # qwen2.5vl 失败的保底：用 opendataloader 结果
            for elem in page_elements.get(pg, []):
                merged_kids.append(elem)
                _count_elem(elem)

    # scanned 页 → qwen2.5vl 结果
    for pg in sorted(scanned_pages):
        if pg in ocr_kids_by_page:
            for elem in ocr_kids_by_page[pg]:
                merged_kids.append(elem)
                _count_elem(elem)

    final_json = {
        "doc_type": "pdf",
        "source": "per-page-adaptive",
        "total_pages": total,
        "total_tables": total_tables,
        "total_paragraphs": total_paras,
        "total_merged_tables": total_merged,
        "total_nested_tables": total_nested,
        "total_diagonal_tables": total_diagonal,
        "kids": merged_kids,
    }

    basename_out = Path(pdf_path).stem
    merged_json_path = Path(output_dir) / f"{basename_out}.json"
    with open(merged_json_path, "w", encoding="utf-8") as f:
        json.dump(final_json, f, ensure_ascii=False, indent=2)
    md_content = convert_pdf_to_markdown_merged(final_json)
    md_path = Path(output_dir) / f"{basename_out}.md"
    with open(md_path, "w", encoding="utf-8") as f:
        f.write(md_content)

    result["success"] = True
    result["mode_used"] = "per-page-adaptive"
    result["files_created"] = [str(merged_json_path), str(md_path)]
    result["classification"] = {
        "total": total,
        "digital": digital_pages,
        "scanned": scanned_pages,
        "digital_with_tables": digital_with_tables,
        "digital_text_only": digital_text_only,
        "scanned_pages": scanned_pages,
        "lang": lang,
    }

    print(f"[成功] 逐页处理完成：{total_tables} 表格（含 {total_merged} 张含合并单元格），{total_paras} 段落")
    print(f"[      ] 纯文本digital页{len(digital_text_only)}页(opendataloader) | "
          f"有表格digital页{len(digital_with_tables)}页(qwen2.5vl) | "
          f"scanned页{len(scanned_pages)}页(qwen2.5vl)")
    return result



def convert_pdf_to_markdown_merged(merged: dict) -> str:
    """从合并结果生成 Markdown"""
    lines = ["# " + merged.get("source", "Document")]
    current_page = None

    for elem in merged.get("kids", []):
        pg = elem.get("page_number")
        if pg != current_page:
            if current_page is not None:
                lines.append("")
            lines.append(f"\n<!-- Page {pg} -->\n")
            current_page = pg

        etype = elem.get("type", "")
        content = elem.get("content", "").strip()
        if not content:
            continue

        if etype == "paragraph":
            lines.append(content)
        elif etype == "table":
            lines.append("")
            lines.append(content)
            lines.append("")

    return "\n".join(lines)


def start_hybrid_server(force_ocr: bool = False, ocr_lang: str = "auto",
                         enrich_formula: bool = False,
                         enrich_picture: bool = False,
                         port: int = HYBRID_SERVER_PORT) -> bool:
    """启动 Hybrid server(后台运行)"""

    # 检查是否已有运行中的 server
    if is_server_running(port):
        print(f"[INFO] Hybrid server 已在运行 (port {port}),跳过启动。")
        return True

    cmd = [JAVA_BIN, "-jar",
           "/home/wangyc/.local/lib/python3.10/site-packages/opendataloader_pdf/jar/opendataloader-pdf-hybrid.jar",
           "--port", str(port)]

    if force_ocr:
        cmd.append("--force-ocr")

    if ocr_lang and ocr_lang != "auto":
        cmd.extend(["--ocr-lang", ocr_lang])

    if enrich_formula:
        cmd.append("--enrich-formula")

    if enrich_picture:
        cmd.append("--enrich-picture-description")

    env = os.environ.copy()
    env["JAVA_HOME"] = JAVA_HOME
    env["PATH"] = JAVA_HOME + "/bin:" + env.get("PATH", "")

    print(f"[INFO] 启动 Hybrid server: {' '.join(cmd)}")

    try:
        proc = subprocess.Popen(
            cmd,
            env=env,
            stdout=subprocess.PIPE,
            stderr=subprocess.STDOUT,
        )

        # 等待 server 启动(最多30秒)
        for _ in range(30):
            time.sleep(1)
            if proc.poll() is not None:
                # 进程已退出
                output = proc.stdout.read().decode(errors="replace")
                print(f"[错误] Hybrid server 启动失败:\n{output[:500]}", file=sys.stderr)
                return False

            if is_server_running(port):
                # 写 PID 到文件
                with open(HYBRID_SERVER_PID_FILE, "w") as f:
                    f.write(str(proc.pid))
                print(f"[INFO] Hybrid server 启动成功 (PID: {proc.pid})")
                return True

        print("[错误] Hybrid server 启动超时", file=sys.stderr)
        proc.kill()
        return False

    except Exception as e:
        print(f"[错误] 启动 Hybrid server 失败: {e}", file=sys.stderr)
        return False


def is_server_running(port: int = HYBRID_SERVER_PORT) -> bool:
    """检查 server 是否在运行"""
    import socket
    try:
        with socket.create_connection(("127.0.0.1", port), timeout=2):
            return True
    except (socket.timeout, ConnectionRefusedError, OSError):
        return False


def stop_hybrid_server():
    """停止 Hybrid server"""
    pid_file = Path(HYBRID_SERVER_PID_FILE)
    if pid_file.exists():
        pid = int(pid_file.read_text().strip())
        try:
            os.kill(pid, 9)
            print(f"[INFO] Hybrid server 已停止 (PID: {pid})")
        except ProcessLookupError:
            print(f"[INFO] Hybrid server 进程已不存在")
        pid_file.unlink()
    else:
        # 尝试通过端口查找
        try:
            import socket
            s = socket.socket()
            s.setsockopt(socket.SOL_SOCKET, socket.SO_REUSEADDR, 1)
            s.bind(("127.0.0.1", HYBRID_SERVER_PORT))
            s.close()
        except OSError:
            # port in use → try kill
            subprocess.run(["fuser", "-k", f"{HYBRID_SERVER_PORT}/tcp"],
                          capture_output=True)


def _post_process_opendataloader_output(pdf_path: str, output_dir: str) -> None:
    """
    opendataloader 输出的 JSON 不含 colspan/row_span。
    本函数对每张表格追加 has_merged_cells 检测(通过列数异常推断),
    并将结果写回 JSON 文件。

    注意:这是"尽力而为"的推断,无法还原精确的 colspan/row_span 值。
    如需精确值,请使用 qwen2.5vl OCR(process_pdf_per_page 路径)。
    """
    import glob
    basename = Path(pdf_path).stem
    json_files = glob.glob(os.path.join(output_dir, f"{basename}*.json"))
    for json_path_str in json_files:
        json_path = Path(json_path_str)
        try:
            with open(json_path, encoding="utf-8") as f:
                data = json.load(f)
        except Exception:
            continue

        modified = False
        for elem in data.get("kids", []):
            if elem.get("type") == "table":
                table_data = elem.get("table_data")
                if table_data and isinstance(table_data, list):
                    flat, has_merged = _detect_merged_cells_from_flat(table_data)
                    if has_merged:
                        elem["has_merged_cells"] = True
                        modified = True
        for elem in data.get("flat_elements", []):
            if elem.get("type") == "table":
                table_data = elem.get("table_data")
                if table_data and isinstance(table_data, list):
                    flat, has_merged = _detect_merged_cells_from_flat(table_data)
                    if has_merged:
                        elem["has_merged_cells"] = True
                        modified = True

        if modified:
            with open(json_path, "w", encoding="utf-8") as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
            print(f"[INFO] 表格合并单元格检测:{json_path.name} 已更新 has_merged_cells 标记")


def build_opendataloader_cmd(pdf_path: str, output_dir: str,
                              detection_info: dict,
                              output_format: str = "markdown,json") -> list:
    """根据检测结果构建 opendataloader-pdf 命令"""

    cmd = [
        sys.executable,  # 当前 Python 解释器
        "-m", "opendataloader_pdf",
        pdf_path,
        "-o", output_dir,
        "-f", output_format,
    ]

    pdf_type = detection_info["type"]

    if pdf_type == "digital":
        # 标准数字 PDF → 本地 Fast 模式,不需要 server
        return [sys.executable, "-m", "opendataloader_pdf",
                pdf_path, "-o", output_dir, "-f", output_format]

    else:
        # 扫描/mixed → Hybrid 模式
        hybrid_cmd = [
            sys.executable, "-m", "opendataloader_pdf",
            "--hybrid", "docling-fast",
        ]

        if pdf_type == "scanned" or detection_info["scanned_pages"] > 0:
            hybrid_cmd.append("--force-ocr")

        hybrid_cmd.extend([pdf_path, "-o", output_dir, "-f", output_format])

        return hybrid_cmd


def run_convert(pdf_path: str, output_dir: str,
                output_format: str = "markdown,json",
                auto_detect: bool = True,
                force_mode: str = None,
                skip_server: bool = False,
                force_qwen: bool = False) -> dict:
    """
    主转换函数

    参数:
        pdf_path: PDF 文件路径
        output_dir: 输出目录
        output_format: 输出格式 (json, markdown, html, text 等)
        auto_detect: 是否自动检测 PDF 类型
        force_mode: 强制模式 ("fast" | "hybrid")
        skip_server: 跳过启动 server(server 已由外部启动)
        force_qwen: 跳过 Hybrid,直接用 qwen2.5vl + Fast(禁用 EasyOCR 作为备选)
    """

    result = {
        "success": False,
        "pdf_path": pdf_path,
        "output_dir": output_dir,
        "detection": None,
        "mode_used": None,
        "command": None,
        "stdout": "",
        "stderr": "",
        "returncode": None,
    }

    # ---- Excel 文档:openpyxl直接提取 + PDF位置辅助 ----
    ext_lower = pdf_path.lower()
    if ext_lower.endswith(".xlsx") or ext_lower.endswith(".xls"):
        print(f"[INFO] 检测到 Excel 文档 → openpyxl 提取 + PDF位置辅助")
        original_basename = Path(pdf_path).stem
        os.makedirs(output_dir, exist_ok=True)

        try:
            # Step 1: openpyxl 提取(结构化数据 + 表格 + 图表)
            print(f"[Step 1/3] Excel 数据提取...")
            excel_data = extract_excel_to_json(pdf_path)

            # Step 2: Excel → PDF(用于页码位置)
            print(f"[Step 2/3] Excel → PDF(用于提取位置信息)...")
            pdf_converted = convert_word_to_pdf(pdf_path, output_dir="/tmp")

            # Step 3: PDF 位置提取
            print(f"[Step 3/3] PDF 位置提取...")
            pdf_json_path = Path("/tmp") / f"{original_basename}.json"
            env = os.environ.copy()
            env["JAVA_HOME"] = JAVA_HOME
            env["PATH"] = JAVA_HOME + "/bin:" + env.get("PATH", "")
            subprocess.run(
                [sys.executable, "-m", "opendataloader_pdf",
                 pdf_converted, "-o", "/tmp", "-f", "json"],
                capture_output=True, text=True, env=env, timeout=300
            )

            # 匹配:按 sheet 名称匹配 page(改进版)
            pdf_page_by_sheet = {}
            pdf_page_count = 0

            if pdf_json_path.exists():
                with open(pdf_json_path, encoding="utf-8") as f:
                    pdf_data = json.load(f)

                pdf_page_count = pdf_data.get("number of pages", 0)

                # 策略1:收集每页的全部文本
                page_texts = {}  # page_num -> all_text
                for elem in pdf_data.get("flat_elements", []):
                    pg = elem.get("page number") or elem.get("page", "?")
                    txt = elem.get("content", "") or ""
                    if pg not in page_texts:
                        page_texts[pg] = ""
                    page_texts[pg] += " " + txt

                # 策略2:建立每页的标题词集合(用于模糊匹配)
                page_headings = {}  # page_num -> set of heading words
                for elem in pdf_data.get("flat_elements", []):
                    if elem.get("type") == "heading":
                        pg = elem.get("page number") or elem.get("page", "?")
                        if pg not in page_headings:
                            page_headings[pg] = set()
                        heading_text = elem.get("content", "") or ""
                        # 拆词
                        for w in re.findall(r"[\w\W]{2,}", heading_text):
                            page_headings[pg].add(w)

                # 为每个 sheet 匹配最可能的 page
                sheets = excel_data.get("sheets", [])

                for sh in sheets:
                    sh_normalized = sh.strip()

                    # 精确匹配:sheet 名出现在某页
                    matched = False
                    for pg, full_text in page_texts.items():
                        if sh_normalized and sh_normalized[:15] in full_text:
                            pdf_page_by_sheet[sh] = pg
                            matched = True
                            break

                    # 模糊匹配:共享词汇多的页面
                    if not matched:
                        sh_words = set(re.findall(r"[\w\W]{2,}", sh_normalized))
                        best_pg, best_score = None, 0
                        for pg, heading_words in page_headings.items():
                            if heading_words and sh_words:
                                score = len(sh_words & heading_words) / len(sh_words | heading_words)
                                if score > best_score and score > 0.1:
                                    best_score = score
                                    best_pg = pg
                        if best_pg:
                            pdf_page_by_sheet[sh] = best_pg
                            matched = True

                    # 未匹配到:不记录(页码留空,人工核对)
                    if not matched:
                        pdf_page_by_sheet[sh] = None

                pdf_json_path.unlink(missing_ok=True)

            # 为每个元素写入 PDF 页码(写入 location["pdf_page"])
            matched_sheets = {}
            for elem in excel_data.get("elements", []):
                sh = elem.get("location", {}).get("excel", "") if isinstance(elem.get("location"), dict) else elem.get("sheet", "")
                if sh in pdf_page_by_sheet and pdf_page_by_sheet[sh]:
                    elem.setdefault("location", {})["pdf_page"] = pdf_page_by_sheet[sh]
                    matched_sheets[sh] = pdf_page_by_sheet[sh]
                else:
                    elem.setdefault("location", {})["pdf_page"] = None

            # 更新 excel_data 元数据
            excel_data["source"] = "openpyxl" + ("+pdf-position" if matched_sheets else "")
            excel_data["sheet_page_map"] = dict(matched_sheets)
            excel_data["pdf_total_pages"] = pdf_page_count
            excel_data["pdf_page_matched"] = len(matched_sheets)
            excel_data["pdf_page_unmatched"] = len(sheets) - len(matched_sheets)

            # 输出 sheet→页码映射供核对(仅显示已匹配到的)
            matched_sheets = {sh: pg for sh, pg in pdf_page_by_sheet.items() if pg}
            if matched_sheets:
                print(f"[INFO] Sheet → PDF页码 映射(共 {pdf_page_count} 页,已匹配 {len(matched_sheets)}/{len(sheets)} 个Sheet):")
                for sh, pg in matched_sheets.items():
                    print(f"       Sheet「{sh}」→ PDF第 {pg} 页")
            if len(matched_sheets) < len(sheets):
                unmatched = [sh for sh in sheets if sh not in matched_sheets]
                print(f"[INFO] 以下 Sheet 未匹配到页码(PDF中未出现Sheet名,请人工核对):")
                for sh in unmatched:
                    print(f"       Sheet「{sh}」→ 需人工确认页码")


            # 保存 JSON
            json_path = Path(output_dir) / f"{original_basename}.json"
            with open(json_path, "w", encoding="utf-8") as f:
                json.dump(excel_data, f, ensure_ascii=False, indent=2)

            # 生成 Markdown
            md_content = extract_excel_to_markdown(pdf_path, page_hints=pdf_page_by_sheet)
            md_path = Path(output_dir) / f"{original_basename}.md"
            with open(md_path, "w", encoding="utf-8") as f:
                f.write(md_content)

            # 清理临时 PDF
            Path(pdf_converted).unlink(missing_ok=True)

            print(f"[成功] Excel 提取完成!")
            print(f"[INFO] 生成文件:")
            print(f"  JSON: {json_path}")
            print(f"  MD:   {md_path}")
            print(f"[INFO] 包含: {excel_data['total_sheets']} 个Sheet | "
                  f"{excel_data['total_tables']} 张表格 | "
                  f"{excel_data['total_charts']} 个图表")

            return {
                "success": True,
                "output_dir": output_dir,
                "mode_used": "excel-openpyxl",
                "files_created": [str(json_path), str(md_path)],
                "excel_stats": {
                    "sheets": excel_data["total_sheets"],
                    "tables": excel_data["total_tables"],
                    "charts": excel_data["total_charts"],
                }
            }

        except Exception as e:
            print(f"[错误] Excel 提取失败: {e}")
            import traceback
            traceback.print_exc()
            return {"success": False, "mode_used": "excel-failed", "error": str(e)}

    # ---- Word 文档:docx内容 + PDF位置 → 合并 ----
    ext_lower = pdf_path.lower()
    if ext_lower.endswith(".docx") or ext_lower.endswith(".doc"):
        print(f"[INFO] 检测到 Word 文档 → 三步处理:docx提取 + PDF位置提取 → 智能合并")
        original_basename = Path(pdf_path).stem
        os.makedirs(output_dir, exist_ok=True)

        env = os.environ.copy()
        env["JAVA_HOME"] = JAVA_HOME
        env["PATH"] = JAVA_HOME + "/bin:" + env.get("PATH", "")

        try:
            # Step 1: Word → PDF
            print(f"[Step 1/3] Word → PDF(用于提取位置信息)...")
            pdf_converted = convert_word_to_pdf(pdf_path, output_dir="/tmp")

            # Step 2: PDF 位置提取(page / bbox / section_path)
            print(f"[Step 2/3] PDF 位置提取...")
            pdf_json_cmd = [
                sys.executable, "-m", "opendataloader_pdf",
                pdf_converted, "-o", "/tmp", "-f", "json",
            ]
            subprocess.run(pdf_json_cmd, capture_output=True, text=True, env=env, timeout=300)

            # opendataloader 输出的文件名就是 basename.json
            pdf_json_path = Path("/tmp") / f"{original_basename}.json"
            pdf_flat = []
            if pdf_json_path.exists():
                with open(pdf_json_path, encoding="utf-8") as f:
                    pdf_data = json.load(f)
                pdf_flat = pdf_data.get("flat_elements", [])
                # 兼容:无 flat_elements 时用 kids 扁平化
                if not pdf_flat and "kids" in pdf_data:
                    pdf_flat = []
                    heading_path = []
                    def _flat(kids, out, hp):
                        for k in kids:
                            if k.get("type") == "heading":
                                level = k.get("heading level", 99)
                                while hp and hp[-1][0] >= level:
                                    hp.pop()
                                hp.append((level, k.get("content", "")))
                                k["section_path"] = " > ".join(x[1] for x in sorted(hp))
                                out.append(k)
                                if "kids" in k:
                                    _flat(k.pop("kids"), out, hp)
                            else:
                                out.append(k)
                                if "kids" in k:
                                    _flat(k.pop("kids"), out, hp)
                    _flat(pdf_data["kids"], pdf_flat, heading_path)
                pdf_json_path.unlink(missing_ok=True)  # 清理临时文件

            print(f"[      ] PDF 提取到 {len(pdf_flat)} 个位置元素")

            # Step 3: docx 内容提取(表格完整 + 文本完整)
            print(f"[Step 3/3] docx 内容提取(完整表格/文本)...")
            docx_data = extract_word_to_json(pdf_path)  # 用已有的 docx 提取函数

            # Step 4: 智能合并
            print(f"[      ] 智能合并:docx内容 + PDF位置 → 统一结构")
            merged = merge_docx_and_pdf(docx_data, pdf_flat)

            # Fallback:合并失败(0元素)时,直接使用 docx 原始内容
            if not merged.get("elements") and not merged.get("kids"):
                print(f"[警告] PDF位置匹配失败(0元素),Fallback:直接使用 docx 内容", file=sys.stderr)
                merged = docx_data  # 直接用 extract_word_to_json 的完整输出
                merged["kids"] = merged.pop("elements", merged.get("kids", []))
                merged["original_word"] = os.path.abspath(pdf_path)

            # 附加元数据
            merged["source"] = "word-to-pdf-merged"
            merged["original_word"] = os.path.abspath(pdf_path)
            merged["converted_pdf"] = pdf_converted

            # 保存合并后的 JSON
            merged_json_path = Path(output_dir) / f"{original_basename}.json"
            with open(merged_json_path, "w", encoding="utf-8") as f:
                json.dump(merged, f, ensure_ascii=False, indent=2)

            # 生成 Markdown(基于 docx 内容,带 section_path)
            md_content = f"# {original_basename}\n\n"
            md_content += convert_word_to_markdown(pdf_path)
            md_path = Path(output_dir) / f"{original_basename}.md"
            with open(md_path, "w", encoding="utf-8") as f:
                f.write(md_content)

            match_stats = merged.get("_pdf_position_hints", {})
            print(f"[成功] 合并完成!")
            print(f"[INFO] 生成文件:")
            print(f"  JSON: {merged_json_path}  (内容:{len(merged.get('kids', []) or merged.get('elements', []))}元素 | 表格:{merged.get('total_tables', 0)} | 段落:{merged.get('total_paragraphs', 0)})")
            print(f"  PDF位置匹配: 表格匹配{match_stats.get('tables_matched','?')}个, 段落匹配{match_stats.get('paras_matched','?')}个")
            print(f"  MD:   {md_path}")
            print(f"[INFO] 临时PDF已清理({pdf_converted})")

            # 清理临时 PDF
            Path(pdf_converted).unlink(missing_ok=True)

            return {
                "success": True,
                "original_word": os.path.abspath(pdf_path),
                "output_dir": output_dir,
                "mode_used": "word-to-pdf-merged",
                "files_created": [str(merged_json_path), str(md_path)],
                "merged_stats": match_stats,
            }

        except Exception as e:
            print(f"[错误] 合并流程失败: {e}")
            import traceback
            traceback.print_exc()
            # 降级:仅 docx 直接提取
            word_result = convert_word(pdf_path, output_dir, output_format)
            word_result["mode_used"] = "word-direct-fallback"
            return word_result
        word_result = convert_word(pdf_path, output_dir, output_format)
        word_result["mode_used"] = "word-direct-fallback"
        return word_result

    # ---- PDF 处理 ----
    # pdf_path 已在参数中传入,继续使用

    # 验证 Java
    if not find_java():
        print("[错误] Java 不可用,请检查 JAVA_HOME 设置", file=sys.stderr)
        return result

    # 解析 PDF 路径
    pdf_path = os.path.abspath(pdf_path)
    if not os.path.exists(pdf_path):
        print(f"[错误] 文件不存在: {pdf_path}", file=sys.stderr)
        return result

    # 自动检测
    if auto_detect and not force_mode:
        print("[INFO] 自动检测 PDF 类型...")
        detection = detect_pdf_type(pdf_path)
        result["detection"] = detection
        print(f"[INFO] 检测结果: {json.dumps(detection, ensure_ascii=False, indent=2)}")
    else:
        detection = {
            "type": "digital" if (force_mode == "fast" or not force_mode) else "scanned",
            "lang": "auto",
            "total_pages": 0,
        }
        result["detection"] = detection

    # 构建命令
    cmd = build_opendataloader_cmd(pdf_path, output_dir, detection, output_format)
    result["command"] = " ".join(cmd)

    print(f"\n[INFO] 使用模式: {'Fast (本地)' if detection['type'] == 'digital' else 'Hybrid (AI增强)'}")
    print(f"[INFO] 执行命令: {' '.join(cmd)}")

    # ============================================================
    # 逐页自适应处理:digital 页用 opendataloader,scanned 页用 OCR
    # ============================================================
    if detection["type"] in ("scanned", "mixed") and not skip_server:
        print(f"\n[INFO] PDF 类型: {detection['type']},启用逐页自适应处理")
        per_page_result = process_pdf_per_page(
            pdf_path=pdf_path,
            output_dir=output_dir,
            output_format=output_format,
            force_qwen=force_qwen,
            skip_server=skip_server,
        )
        if per_page_result.get("success"):
            result["success"] = True
            result["mode_used"] = per_page_result.get("mode_used", "per-page-adaptive")
            result["files_created"] = per_page_result.get("files_created", [])
            result["classification"] = per_page_result.get("classification", {})
            return result
        else:
            # 逐页处理失败,降级到传统 Fast 模式
            print("[警告] 逐页处理失败,降级到 Fast 本地模式...", file=sys.stderr)
            detection["type"] = "digital"
            cmd = build_opendataloader_cmd(pdf_path, output_dir, detection, output_format)
            result["command"] = " ".join(cmd)
    else:
        # digital PDF:直接构建命令
        cmd = build_opendataloader_cmd(pdf_path, output_dir, detection, output_format)
        result["command"] = " ".join(cmd)

    # 执行
    env = os.environ.copy()
    env["JAVA_HOME"] = JAVA_HOME
    env["PATH"] = JAVA_HOME + "/bin:" + env.get("PATH", "")

    try:
        proc = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            env=env,
            timeout=300  # 5分钟超时
        )
        result["stdout"] = proc.stdout
        result["stderr"] = proc.stderr
        result["returncode"] = proc.returncode

        if proc.returncode == 0:
            result["success"] = True
            result["mode_used"] = detection["type"]
            print(f"[成功] 输出已保存到: {output_dir}")
            # Post-process: opendataloader 输出的 JSON 中,表格追加 has_merged_cells 检测
            _post_process_opendataloader_output(pdf_path, output_dir)
        else:
            print(f"[错误] 转换失败 (返回码 {proc.returncode}):", file=sys.stderr)
            print(proc.stderr[:500] if proc.stderr else proc.stdout[:500], file=sys.stderr)

    except subprocess.TimeoutExpired:
        result["stderr"] = "超时(5分钟)"
        print("[错误] 转换超时(5分钟)", file=sys.stderr)
    except Exception as e:
        result["stderr"] = str(e)
        print(f"[错误] {e}", file=sys.stderr)

    return result


# ---------- CLI 入口 ----------
def main():
    parser = argparse.ArgumentParser(
        description="PDF/Word 文档自动检测与转换工具(统一接口)"
    )
    parser.add_argument("input", help="输入 PDF 文件路径")
    parser.add_argument("-o", "--output", required=True, help="输出目录")
    parser.add_argument("-f", "--format",
                        default="markdown,json",
                        help="输出格式(逗号分隔,默认 markdown,json)")
    parser.add_argument("--no-auto-detect", dest="auto_detect",
                        action="store_false", default=True,
                        help="禁用自动检测")
    parser.add_argument("--force-mode",
                        choices=["fast", "hybrid"],
                        default=None,
                        help="强制指定模式")
    parser.add_argument("--skip-server", action="store_true",
                        help="跳过 server 启动(server 已运行时使用)")
    parser.add_argument("--stop-server", action="store_true",
                        help="停止 Hybrid server")
    parser.add_argument("--detect-only", action="store_true",
                        help="仅检测 PDF 类型,不执行转换")
    parser.add_argument("--force-qwen", action="store_true",
                        help="跳过 Hybrid,直接用 qwen2.5vl + Fast(禁用 EasyOCR 作为备选)")

    args = parser.parse_args()

    # 处理 stop-server
    if args.stop_server:
        stop_hybrid_server()
        return

    # 处理 detect-only
    if args.detect_only:
        ext = args.input.lower()
        if ext.endswith(".docx") or ext.endswith(".doc"):
            info = {"type": "word", "format": "docx/doc"}
        elif ext.endswith(".xlsx") or ext.endswith(".xls"):
            info = {"type": "excel", "format": "xlsx/xls"}
        else:
            info = detect_pdf_type(args.input)
        print(json.dumps(info, ensure_ascii=False, indent=2))
        return

    # 创建输出目录
    os.makedirs(args.output, exist_ok=True)

    # 执行转换
    result = run_convert(
        pdf_path=args.input,
        output_dir=args.output,
        output_format=args.format,
        auto_detect=args.auto_detect,
        force_mode=args.force_mode,
        skip_server=args.skip_server,
        force_qwen=args.force_qwen,
    )

    # 输出结果 JSON
    print("\n--- Result ---")
    print(json.dumps({
        "success": result["success"],
        "mode": result["mode_used"],
        "detection": result.get("detection"),
        "files_created": result.get("files_created", []),
        "command": result.get("command"),
    }, ensure_ascii=False, indent=2))

    sys.exit(0 if result["success"] else 1)


if __name__ == "__main__":
    main()
